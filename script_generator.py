"""
AI 选题与剧本生成模块
用 LLM（OpenAI 兼容接口）完成两步：
1. 热点 → 叙事短片选题（Step 2 + Step 3）
2. 选题 → 三幕剧分镜剧本（Step 4）

面向 B站叙事 AI 短片赛道，剧本用于可灵（Kling）AI 多镜头视频生成：
有主角、有冲突、有反转，分场分镜，每个镜头带可灵视觉提示。
输出结构化 JSON，方便后续角色设定卡 / 场景九宫格 / 可灵镜头生成模块直接消费。
"""
import json
import re
from typing import List, Dict, Any, Union
import config


# ===================== JSON 健壮解析器 =====================

def _robust_json_loads(raw: str, expect_type: type = None) -> Union[List, Dict]:
    """
    多策略 JSON 解析器，专治 LLM 输出的"不标准 JSON"。
    按顺序尝试 5 种修复策略，全部失败才抛异常。

    常见 LLM JSON 问题：
    1. 包 ```json ... ``` 代码块
    2. 前后带解释文字
    3. 字符串值里含未转义的双引号
    4. 字符串值里含未转义的换行
    5. 尾部多余的逗号
    6. 单引号代替双引号
    """
    raw = raw.strip() if raw else ""

    # 策略1：去掉 markdown 代码块标记
    if raw.startswith("```"):
        # 去掉首行 ``` 或 ```json
        lines = raw.split("\n")
        if lines and lines[0].startswith("```"):
            lines = lines[1:]
        # 去掉末尾 ```
        if lines and lines[-1].strip() == "```":
            lines = lines[:-1]
        raw = "\n".join(lines).strip()

    # 策略2：直接尝试解析
    try:
        result = json.loads(raw)
        if _check_type(result, expect_type):
            return result
    except json.JSONDecodeError:
        pass

    # 策略3：提取首个 [ 到末尾 ] 或首个 { 到末尾 }
    open_char = "[" if expect_type == list else "{"
    close_char = "]" if expect_type == list else "}"
    # 如果没指定类型，优先尝试从第一个 { 或 [ 开始
    if expect_type is None:
        first_obj = raw.find("{")
        first_arr = raw.find("[")
        if first_obj >= 0 and (first_arr < 0 or first_obj < first_arr):
            open_char, close_char = "{", "}"
        elif first_arr >= 0:
            open_char, close_char = "[", "]"

    start = raw.find(open_char)
    end = raw.rfind(close_char)
    if start >= 0 and end > start:
        candidate = raw[start : end + 1]
        try:
            result = json.loads(candidate)
            if _check_type(result, expect_type):
                return result
        except json.JSONDecodeError:
            pass

        # 策略4：修复常见 JSON 格式问题后重试
        fixed = _fix_common_json_issues(candidate)
        try:
            result = json.loads(fixed)
            if _check_type(result, expect_type):
                return result
        except json.JSONDecodeError:
            pass

        # 策略5：用正则逐个提取键值对（最后兜底，针对 dict）
        if expect_type == dict or (expect_type is None and open_char == "{"):
            result = _regex_extract_dict(candidate)
            if result and _check_type(result, expect_type):
                return result

    # 全部失败，抛出带原始内容的异常，方便调试
    raise json.JSONDecodeError(
        f"无法解析 LLM 输出为 JSON。原始内容前 300 字符：\n{raw[:300]}",
        raw, 0
    )


def _check_type(result, expect_type) -> bool:
    """检查解析结果类型是否符合预期。"""
    if expect_type is None:
        return True
    if expect_type == list:
        return isinstance(result, list)
    if expect_type == dict:
        return isinstance(result, dict)
    return True


def _fix_common_json_issues(s: str) -> str:
    """修复 LLM 输出 JSON 时的常见格式问题。"""
    # 1. 去掉尾部多余的逗号（}, ] 前的逗号）
    s = re.sub(r",\s*([}\]])", r"\1", s)

    # 2. 单引号转双引号（注意不能误伤已经是双引号的内容）
    # 这个操作比较危险，只在纯单引号场景下做
    if "'" in s and '"' not in s:
        s = s.replace("'", '"')

    # 3. 修复字符串值内未转义的换行符
    # 找到所有双引号字符串，把内部的换行替换成 \n
    def fix_string_newlines(match):
        content = match.group(1)
        # 把实际换行替换成转义换行
        content = content.replace("\n", "\\n").replace("\r", "\\r")
        return f'"{content}"'
    # 匹配 "..." 但不匹配已经被转义的
    s = re.sub(r'"((?:[^"\\]|\\.)*?)(?<!\\)"', fix_string_newlines, s, flags=re.DOTALL)

    # 4. 修复键名没加引号的情况 {title: "..."} → {"title": "..."}
    s = re.sub(r'(\{|,)\s*([a-zA-Z_][a-zA-Z0-9_]*)\s*:', r'\1"\2":', s)

    # 5. 把中文引号统一成英文双引号
    s = s.replace("\u201c", '"').replace("\u201d", '"')
    s = s.replace("\u2018", '"').replace("\u2019", '"')

    return s


def _regex_extract_dict(s: str) -> Dict:
    """最后兜底：用正则逐个提取顶层键值对，构造 dict。"""
    result = {}
    # 匹配 "key": "value" 或 "key": number 或 "key": [...] 或 "key": {...}
    # 这个方法比较粗糙，但能救回一些极端情况
    pattern = r'"([^"]+)"\s*:\s*("(?:[^"\\]|\\.)*"|[\d.]+|\[[\s\S]*?\]|\{[\s\S]*?\}|true|false|null)'
    for match in re.finditer(pattern, s):
        key = match.group(1)
        val_str = match.group(2).strip()
        # 尝试解析值
        try:
            val = json.loads(val_str)
        except json.JSONDecodeError:
            # 去掉首尾引号当作字符串
            if val_str.startswith('"') and val_str.endswith('"'):
                val = val_str[1:-1]
            else:
                val = val_str
        # 同名 key 只取第一个
        if key not in result:
            result[key] = val
    return result


# ===================== LLM 调用封装 =====================

try:
    from openai import OpenAI
    HAS_OPENAI = True
except ImportError:
    HAS_OPENAI = False
    OpenAI = None


def _sanitize_ascii(value: str, field_name: str = "API Key") -> str:
    """
    清除字符串中的非 ASCII 字符。
    用户从 Word/PDF/网页复制 API Key 时，可能混入不可见的特殊字符
    （如 U+F0B7 Symbol 字体的项目符号、零宽空格 U+200B、BOM U+FEFF 等），
    这些字符会导致 httpx 构造 HTTP 头时报 UnicodeEncodeError。
    """
    if not value:
        return value
    # 去掉首尾空白（包括不可见空白字符）
    cleaned = value.strip()
    # 去掉零宽字符、BOM、软连字符等不可见字符
    invisible_chars = {
        "\ufeff",  # BOM / Zero Width No-Break Space
        "\u200b",  # Zero Width Space
        "\u200c",  # Zero Width Non-Joiner
        "\u200d",  # Zero Width Joiner
        "\u200e",  # Left-to-Right Mark
        "\u200f",  # Right-to-Left Mark
        "\u00ad",  # Soft Hyphen
    }
    for ch in invisible_chars:
        cleaned = cleaned.replace(ch, "")
    # 只保留 ASCII 可见字符（0x20-0x7E）
    ascii_only = "".join(c for c in cleaned if 0x20 <= ord(c) <= 0x7E)
    if len(ascii_only) != len(cleaned):
        removed = len(cleaned) - len(ascii_only)
        print(f"[配置] ⚠️ {field_name} 含 {removed} 个非 ASCII 字符，已自动清除")
    return ascii_only


def _get_client():
    if not HAS_OPENAI:
        raise RuntimeError("未安装 openai 包，请 pip install openai")
    if not config.LLM_API_KEY:
        raise RuntimeError("LLM_API_KEY 未配置，请检查 .env")
    # 清除 API Key 和 Base URL 中可能混入的非 ASCII 字符
    clean_key = _sanitize_ascii(config.LLM_API_KEY, "LLM_API_KEY")
    clean_base = _sanitize_ascii(config.LLM_API_BASE, "LLM_API_BASE")
    if not clean_key:
        raise RuntimeError("LLM_API_KEY 清除非 ASCII 字符后为空，请检查是否复制完整")
    return OpenAI(api_key=clean_key, base_url=clean_base)


def _chat(system: str, user: str, temperature: float = 0.8) -> str:
    """统一的 chat 调用封装。"""
    client = _get_client()
    resp = client.chat.completions.create(
        model=config.LLM_MODEL,
        messages=[
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
        temperature=temperature,
    )
    return resp.choices[0].message.content


# ===================== Step 2+3: 热点 → 选题 =====================

def generate_topics(hot_topics: List[Dict[str, Any]], count: int = None) -> List[Dict[str, Any]]:
    """
    输入热点列表，输出可影视化的叙事短片选题。
    把当日热点"故事化"——从热点里提炼出有主角、有冲突、有反转的短片构思，
    而非平铺直叙的知识讲解。
    每个选题包含：title, hook, angle, genre, premise, hot_source, science_keywords, duration_estimate
    """
    count = count or config.TOPIC_COUNT

    # 只取标题给 LLM，避免 token 浪费
    hot_titles = [t["title"] for t in hot_topics[:30]]
    genre_options = "、".join(config.GENRE_OPTIONS)

    system = f"""你是一位专注于「{config.NICHE}」赛道的 B站叙事短片选题策划师。
你的任务：从给定的当日热点中，提炼出能改编成 3-5 分钟叙事短片的故事构思。
不是做知识科普讲解，而是把热点"故事化"——找出一个有主角、有冲突、有反转的切入点。

举例：
- 热点"深海发现新物种发光水母" → 短片《深海来客》：一个生物学家在深渊发现了不该被发现的东西
- 热点"AI 通过图灵测试" → 短片《最后一题》：面试官发现对面的候选人不是人类
- 热点"火星发现液态水痕迹" → 短片《倒数第七天》：驻火星基地的工程师收到一条不该存在的水源信号

选题原则：
1. 必须能化作一个"有人物、有冲突"的小故事，不是知识罗列
2. 必须有反转空间——观众看完会有"原来如此"的意外感
3. 时长可控在 3-5 分钟，不要宏大叙事
4. 题材类型从以下选一个：{genre_options}
5. 视觉上可由可灵 AI 生成（角色 + 场景），不依赖真实新闻画面

输出严格的 JSON 数组，每个元素包含：
- title: 短片标题（15字以内，带悬念，有电影感）
- hook: 开场钩子（前3秒画面/旁白话术，必须抓人）
- angle: 叙事切入角度（一句话说明从哪个故事点讲起、主角是谁）
- genre: 题材类型（从 {genre_options} 中选一个）
- premise: 故事一句话梗概（主角 + 冲突 + 反转方向，不要剧透结局）
- hot_source: 对应的原始热点标题
- science_keywords: 3-5 个视觉/故事关键词（用于后续生成角色与场景视觉参考，如 ["深海","潜水艇","发光生物","幽闭空间"]）
- duration_estimate: 预计时长（秒）

只输出 JSON，不要任何解释文字。"""

    user = f"今日热点列表：\n" + "\n".join(f"{i+1}. {t}" for i, t in enumerate(hot_titles)) + \
           f"\n\n请从中挑选 {count} 个最适合改编成叙事短片的，按上述格式输出 JSON 数组。"

    raw = _chat(system, user, temperature=0.85)
    topics = _robust_json_loads(raw, expect_type=list)

    print(f"[选题] 生成 {len(topics)} 个叙事短片选题")
    for t in topics:
        print(f"  - {t.get('title')}（{t.get('genre')} | 关键词: {t.get('science_keywords')}）")
    return topics


# ===================== Step 4: 选题 → 分镜脚本 =====================

def generate_script(topic: Dict[str, Any]) -> Dict[str, Any]:
    """
    输入叙事短片选题，输出完整的 B站叙事短片剧本（三幕剧结构）。

    用于可灵（Kling）AI 多镜头视频生成：按场景分组连续生成以保持角色与场景一致性。
    剧本包含主角、冲突、反转，分场分镜，每个镜头带可灵视觉提示。
    旁白/对话字数严格按 字数 ≈ duration × NARRATION_CHARS_PER_SEC 匹配，
    避免文案与视频时长错位。

    返回结构：
    {
      "title": "短片标题",
      "genre": "科幻/悬疑/哲理/奇幻/都市/恐怖",
      "theme": "主题一句话",
      "total_duration": 240,
      "narration_full": "完整旁白/对话（拼接所有镜头）",
      "characters": [
        {
          "name": "角色名",
          "role": "主角/配角",
          "appearance": "外貌描述（性别、年龄、发型、服装、特征）",
          "personality": "性格描述",
          "voice": "声音特征"
        }
      ],
      "scenes": [
        {
          "id": 1,
          "name": "场景名称",
          "location": "地点描述",
          "time_of_day": "日/夜/黄昏",
          "weather": "天气/氛围",
          "description": "场景视觉描述（供生成场景九宫格用）"
        }
      ],
      "shots": [
        {
          "index": 1,
          "scene_id": 1,
          "scene_name": "场景名称",
          "duration": 8,
          "narration": "约32字",
          "visual_description": "画面描述",
          "camera_movement": "推/拉/摇/移/跟/固定",
          "character_action": "角色动作",
          "mood": "情绪"
        }
      ],
      "bgm_mood": "mysterious"
    }
    """
    target = config.TARGET_DURATION
    cps = config.NARRATION_CHARS_PER_SEC  # 字/秒
    min_shots = config.MIN_SHOTS
    max_shots = config.MAX_SHOTS
    genre = topic.get("genre") or config.DEFAULT_GENRE
    genre_options = "、".join(config.GENRE_OPTIONS)

    system = f"""你是一位 B站叙事 AI 短片编剧，专攻「{config.NICHE}」赛道，用可灵（Kling）AI 生成视频。
你要把一个选题扩写成一部 3-5 分钟的叙事短片完整剧本，要求：有主角、有冲突、有反转，三幕剧结构。

【题材类型】
当前选题题材倾向：{genre}（可选类型：{genre_options}）。剧本基调需契合题材。

【三幕剧结构（总时长约 {target} 秒，可浮动 ±30 秒）】
- 第一幕·建置（约 25-30% 时长）：交代主角、世界观、日常，结尾抛出"激励事件"打破平衡
- 第二幕·对抗（约 45-55% 时长）：主角陷入困境/追查/冲突升级，中点出现一次小反转，结尾达到危机高潮
- 第三幕·结局（约 20-25% 时长）：意料之外的主反转揭示真相，主角做出选择，余味收尾
绝对不要平铺直叙地讲解知识，必须有明确的矛盾冲突和至少一个意料之外的转折。

【主角设定】
1-2 个主要角色，每个角色必须有：名字、性格、外貌（性别/年龄/发型/服装/特征）、声音特征。
外貌描述要具体到可灵能据此生成"角色设定卡"（一致性参考）的程度。

【分场分镜——为可灵多镜头连续生成设计】
1. 全片拆成 3-6 个场景（scene），同一场景内的镜头连续生成以保持角色与场景一致性
2. 每个场景内 3-6 个镜头（shot），全片镜头总数 {min_shots}-{max_shots} 个
3. 场景需给出：名称、地点、时段（日/夜/黄昏）、天气/氛围、视觉描述（供生成场景九宫格）

【旁白/对话——严格匹配时长，这是硬性要求】
中文语速约 {cps} 字/秒。每个镜头的 narration 字数必须严格等于 duration × {cps}（允许 ±2 字误差）。
公式：narration 字数 = duration × {cps}
举例：duration=8 秒 → narration 约 {8*cps} 字；duration=15 秒 → narration 约 {15*cps} 字
narration 可以是旁白、角色独白或对白，按叙事需要混用，绝对不能出现文案与时长错位的情况。

【每个镜头要求】
1. index: 镜头序号（从 1 递增）
2. scene_id: 所属场景 ID（与 scenes[].id 对应，同一场景的镜头 scene_id 相同且 index 连续）
3. scene_name: 场景名称（与 scenes[].name 一致）
4. duration: 时长（秒），整数
5. narration: 旁白/对白（字数严格 = duration × {cps}）
6. visual_description: 画面描述（给可灵的视觉提示，要具体：人物动作、表情、环境光影、构图、色调）
7. camera_movement: 镜头运动，从"推/拉/摇/移/跟/固定"中选一个
8. character_action: 角色动作描述（供可灵主体参考用，无角色时填"无"）
9. mood: 情绪氛围（如 紧张/温馨/神秘/震撼/压抑/释然）

输出严格 JSON，结构如下：
{{
  "title": "短片标题",
  "genre": "{genre}",
  "theme": "主题一句话",
  "total_duration": 实际总时长,
  "narration_full": "所有镜头 narration 拼接的完整旁白/对话",
  "characters": [
    {{"name": "角色名", "role": "主角", "appearance": "外貌描述", "personality": "性格描述", "voice": "声音特征"}}
  ],
  "scenes": [
    {{"id": 1, "name": "场景名称", "location": "地点", "time_of_day": "日/夜/黄昏", "weather": "天气/氛围", "description": "场景视觉描述"}}
  ],
  "shots": [
    {{"index": 1, "scene_id": 1, "scene_name": "场景名称", "duration": 8, "narration": "约32字", "visual_description": "画面描述", "camera_movement": "推", "character_action": "角色动作", "mood": "情绪"}}
  ],
  "bgm_mood": "mysterious|tense|epic|melancholic|warm|dark"
}}

只输出 JSON，不要解释。"""

    user = f"""请为以下选题创作叙事短片剧本：

选题标题：{topic.get('title')}
开场钩子：{topic.get('hook')}
叙事切入角度：{topic.get('angle')}
题材倾向：{genre}
故事梗概参考：{topic.get('premise', '无')}
对应热点：{topic.get('hot_source', '无')}
视觉/故事关键词参考：{topic.get('science_keywords', topic.get('story_keywords', []))}

要求：
1. 三幕剧结构，总时长接近 {target} 秒（可浮动 ±30 秒）
2. 1-2 个有名字、有性格、有外貌的主角
3. 明确的矛盾冲突 + 至少一个意料之外的反转
4. 分 3-6 个场景，每场景 3-6 个镜头，全片 {min_shots}-{max_shots} 个镜头
5. 同一场景的镜头 scene_id 相同且 index 连续（方便可灵连续生成保持一致性）
6. 每个镜头的 narration 字数严格等于 duration × {cps}（±2 字）"""

    raw = _chat(system, user, temperature=0.8)
    script = _robust_json_loads(raw, expect_type=dict)

    # 后处理：校验并修正旁白/对白字数
    _validate_and_fix_narration_length(script, cps)

    print(f"[剧本] 生成 {len(script.get('shots', []))} 个镜头 / "
          f"{len(script.get('scenes', []))} 个场景 / "
          f"{len(script.get('characters', []))} 个角色，总时长 {script.get('total_duration')}s")
    return script


def _validate_and_fix_narration_length(script: Dict[str, Any], cps: int):
    """
    校验每个分镜的解说文案字数是否匹配时长。
    不匹配时打印警告（不强制截断，保留 AI 原文，但标记差异供后续调整）。
    """
    shots = script.get("shots", [])
    issues = []
    for shot in shots:
        duration = shot.get("duration", 0)
        narration = shot.get("narration", "")
        actual_len = len(narration)
        target_len = duration * cps
        if abs(actual_len - target_len) > 5:
            issues.append(
                f"  分镜{shot.get('index')}: 时长{duration}s 需约{target_len}字，实际{actual_len}字（差{actual_len-target_len:+d}）"
            )
            # 标记到分镜中，供 adjust 阶段参考
            shot["narration_length_issue"] = {
                "target": target_len,
                "actual": actual_len,
                "diff": actual_len - target_len,
            }

    if issues:
        print(f"[脚本] ⚠️ {len(issues)} 个分镜解说字数与时长不匹配：")
        for issue in issues:
            print(issue)
    else:
        print(f"[脚本] ✅ 所有分镜解说字数与时长匹配")


# ===================== Step 4.6: 根据实际素材调整脚本 =====================

def adjust_script_to_clips(
    script: Dict[str, Any],
    clips_map: Dict[int, List[Dict[str, Any]]],
) -> Dict[str, Any]:
    """
    根据实际下载的视频素材，动态调整脚本时长和解说建议。

    参数：
    - script: generate_script() 的输出
    - clips_map: {shot_index: [clip_info, ...]}，clip_info 含 duration 字段

    返回调整后的 script（新增 actual_duration, duration_adjustment, adjust_tip 字段）。
    """
    cps = config.NARRATION_CHARS_PER_SEC
    shots = script.get("shots", [])
    adjusted_shots = []
    total_adjusted = 0

    for shot in shots:
        idx = shot.get("index")
        planned_duration = shot.get("duration", 0)
        planned_narration = shot.get("narration", "")
        narration_len = len(planned_narration)

        clips = clips_map.get(idx, [])
        clip_duration = clips[0].get("duration", 0) if clips else 0

        if not clips:
            # 没抓到素材，保留原计划
            shot["actual_duration"] = planned_duration
            shot["duration_adjustment"] = "no_clip"
            shot["adjust_tip"] = (
                f"⚠️ 未抓到素材，需手动找视频。计划时长{planned_duration}秒，"
                f"解说{narration_len}字（约{narration_len/cps:.1f}秒）。"
            )
            total_adjusted += planned_duration
            adjusted_shots.append(shot)
            continue

        # 计算解说需要的时长
        narration_duration = narration_len / cps

        # 策略：以解说时长为准，调整素材使用方式
        if clip_duration >= narration_duration:
            # 素材够长，截取解说时长对应的片段
            actual_duration = round(narration_duration)
            trim = round(clip_duration - narration_duration)
            if trim > 2:
                adjust_tip = (
                    f"✅ 素材{clip_duration}秒 > 解说需{narration_duration:.1f}秒。"
                    f"截取素材精华{actual_duration}秒，剪掉末尾{trim}秒冗余。"
                    f"建议取素材中最精彩的{actual_duration}秒段落。"
                )
            else:
                adjust_tip = (
                    f"✅ 素材{clip_duration}秒 ≈ 解说需{narration_duration:.1f}秒，"
                    f"直接使用，无需裁剪。"
                )
            shot["actual_duration"] = actual_duration
            shot["duration_adjustment"] = f"trim_{trim}s"
        else:
            # 素材不够长，需要处理
            gap = round(narration_duration - clip_duration)
            if gap <= 3:
                # 差距小，素材可适当放慢或重复片段
                adjust_tip = (
                    f"⚠️ 素材{clip_duration}秒 < 解说需{narration_duration:.1f}秒（差{gap}秒）。"
                    f"建议：素材末尾做0.5倍速慢放{gap}秒，或循环播放最后2秒填补。"
                )
                shot["actual_duration"] = round(narration_duration)
                shot["duration_adjustment"] = f"extend_{gap}s"
            else:
                # 差距大，建议精简解说
                target_narration_len = int(clip_duration * cps)
                adjust_tip = (
                    f"⚠️ 素材仅{clip_duration}秒，但解说需{narration_duration:.1f}秒（差{gap}秒）。"
                    f"建议方案A：精简解说至约{target_narration_len}字（删减{narration_len-target_narration_len}字）。"
                    f"建议方案B：补1-2个B-roll空镜填补{gap}秒。"
                )
                shot["actual_duration"] = clip_duration
                shot["duration_adjustment"] = f"shorten_narration_{gap}s"

        shot["adjust_tip"] = adjust_tip
        shot["clip_duration"] = clip_duration
        shot["clip_source"] = clips[0].get("source", "")
        total_adjusted += shot["actual_duration"]
        adjusted_shots.append(shot)

    # 更新脚本
    script["shots"] = adjusted_shots
    script["actual_total_duration"] = total_adjusted
    script["planned_total_duration"] = script.get("total_duration", 0)
    duration_diff = total_adjusted - script.get("total_duration", 0)
    script["duration_diff"] = duration_diff

    print(f"[调整] 计划总时长 {script.get('total_duration',0)}s → 实际总时长 {total_adjusted}s（{duration_diff:+d}s）")
    return script


# ===================== Step 4.5: 生成发布信息 =====================

def generate_publish_info(topic: Dict[str, Any], script: Dict[str, Any]) -> Dict[str, Any]:
    """
    基于选题和剧本，生成 B站发布所需的作品信息。
    返回结构：
    {
      "title": "≤30字的B站标题",
      "description": "≤300字的B站简介（可分段）",
      "hashtags": ["#标签1", ..., "#标签10"],
      "partition": "B站分区推荐（如 科技-数码 / 动画-短片 / 知识-科普）",
      "cover_suggestion": "封面文字建议",
      "publish_tips": "发布时机和注意事项建议"
    }
    """
    system = f"""你是一位 B站叙事短片赛道的运营专家，精通 B站算法和用户心理。
你要为一部已经写好的叙事 AI 短片生成发布所需的文案。

B站发布文案规则：
1. 作品标题（≤30字）：可以有标题党，但必须保留信息量；带悬念/反转钩子，不直接剧透结局；契合 B站社区调性
2. 作品简介（≤300字）：可分段；首句钩子，中间交代短片世界观/看点，结尾引导三连/弹幕互动
3. 标签（≤10个）：1-2个分区大词（如 #短片 #AI短片）+ 3-5个题材/主题相关词 + 1-2个蹭热点词；要真实存在且有热度
4. 分区推荐：根据题材从 B站常见分区中选一个最合适的（如 科技-数码 / 动画-短片 / 知识-科普 / 影视-短片 等）
5. 不要用"震惊""不转不是X"等低质诱导词

输出严格 JSON：
{{
  "title": "作品标题，≤30字",
  "description": "作品简介，≤300字",
  "hashtags": ["#标签1", "#标签2", "..."],
  "partition": "B站分区推荐",
  "cover_suggestion": "封面应突出的关键词或文字（10字以内）",
  "publish_tips": "发布时机和注意事项（一句话）"
}}

只输出 JSON，不要解释。"""

    user = f"""请为以下叙事短片生成 B站发布文案：

短片标题：{script.get('title', topic.get('title'))}
题材：{script.get('genre', topic.get('genre', '未知'))}
主题：{script.get('theme', '')}
对应热点：{topic.get('hot_source', '无')}
BGM情绪：{script.get('bgm_mood', 'mysterious')}
视频时长：{script.get('total_duration', 240)}秒
角色：{', '.join(c.get('name', '') for c in script.get('characters', []))}

旁白/对白摘要（前150字）：
{script.get('narration_full', '')[:150]}

要求：
- 标题≤30字，带悬念但保留信息量
- 简介≤300字，可分段，首句钩子+结尾互动引导
- 最多10个标签，含分区大词+题材词+蹭热点词
- 给出最合适的 B站分区推荐
- 整体风格符合 B站叙事短片社区调性"""

    raw = _chat(system, user, temperature=0.85)
    info = _robust_json_loads(raw, expect_type=dict)

    # 安全校验：超长截断，标签数限制
    info["title"] = info.get("title", "")[:30]
    info["description"] = info.get("description", "")[:300]
    info["hashtags"] = info.get("hashtags", [])[:10]
    info.setdefault("partition", "")
    info.setdefault("cover_suggestion", "")
    info.setdefault("publish_tips", "")

    print(f"[发布] 标题({len(info['title'])}字): {info['title']}")
    print(f"[发布] 简介({len(info['description'])}字)")
    print(f"[发布] 标签({len(info['hashtags'])}个): {' '.join(info['hashtags'])}")
    print(f"[发布] 分区: {info.get('partition')}")
    return info


def publish_info_to_markdown(info: Dict[str, Any]) -> str:
    """把发布信息转成易读的 Markdown 文档（B站风格）。"""
    lines = [
        "# 📱 B站发布文案（直接复制使用）",
        "",
        "## 作品标题",
        "",
        f"**{info.get('title', '')}**",
        "",
        "## 作品简介",
        "",
        info.get("description", ""),
        "",
        "## 标签（发布时填到标签栏）",
        "",
        " ".join(info.get("hashtags", [])),
        "",
        "## 分区推荐",
        "",
        f"**{info.get('partition', '')}**",
        "",
        "## 封面建议",
        "",
        f"封面文字突出：**{info.get('cover_suggestion', '')}**",
        "",
        "## 发布建议",
        "",
        info.get("publish_tips", ""),
        "",
        "---",
        "",
        "## 📋 复制清单",
        "",
        "### 标题（复制到标题栏）",
        "```",
        info.get("title", ""),
        "```",
        "",
        "### 简介（复制到简介栏）",
        "```",
        info.get("description", ""),
        "```",
        "",
        "### 标签（复制到标签栏）",
        "```",
        " ".join(info.get("hashtags", [])),
        "```",
    ]
    return "\n".join(lines)


def script_to_markdown(topic: Dict[str, Any], script: Dict[str, Any]) -> str:
    """把叙事短片剧本转成易读的 Markdown 文档，方便用可灵 AI 生成视频时对照。"""
    lines = [
        f"# {script.get('title', topic.get('title', '未命名短片'))}",
        "",
        f"**赛道**：{config.NICHE}",
        f"**题材**：{script.get('genre', topic.get('genre', '未知'))}",
        f"**主题**：{script.get('theme', '')}",
        f"**对应热点**：{topic.get('hot_source', '无')}",
        f"**预计时长**：{script.get('total_duration', '?')} 秒",
    ]
    # 如果有调整后的实际时长，也显示
    if script.get("actual_total_duration"):
        lines.append(f"**实际时长（根据素材调整）**：{script.get('actual_total_duration')} 秒（原计划 {script.get('planned_total_duration','?')} 秒，差 {script.get('duration_diff', 0):+d} 秒）")
    lines.extend([
        f"**BGM 情绪标签**：{script.get('bgm_mood', 'mysterious')}",
        "",
        "## 角色设定",
        "",
    ])
    for c in script.get("characters", []):
        lines.append(f"### {c.get('name', '未命名')}（{c.get('role', '角色')}）")
        lines.append(f"- **外貌**：{c.get('appearance', '')}")
        lines.append(f"- **性格**：{c.get('personality', '')}")
        lines.append(f"- **声音**：{c.get('voice', '')}")
        lines.append("")
    lines.extend([
        "## 场景设定",
        "",
    ])
    for s in script.get("scenes", []):
        lines.append(f"### 场景{s.get('id', '?')}：{s.get('name', '未命名')}")
        lines.append(f"- **地点**：{s.get('location', '')}")
        lines.append(f"- **时段**：{s.get('time_of_day', '')}")
        lines.append(f"- **天气/氛围**：{s.get('weather', '')}")
        lines.append(f"- **视觉描述**：{s.get('description', '')}")
        lines.append("")
    lines.extend([
        "## 完整旁白/对话（可直接复制做 AI 配音）",
        "",
        script.get("narration_full", ""),
        "",
        "## 分镜表",
        "",
        "| 序号 | 场景 | 时长 | 旁白/对话 | 画面描述 | 镜头运动 | 角色动作 | 情绪 |",
        "|------|------|------|-----------|----------|----------|----------|------|",
    ])
    for shot in script.get("shots", []):
        narration = shot.get("narration", "").replace("|", "/").replace("\n", " ")
        visual = shot.get("visual_description", "").replace("|", "/").replace("\n", " ")
        action = shot.get("character_action", "").replace("|", "/").replace("\n", " ")
        scene_name = str(shot.get("scene_name", "")).replace("|", "/").replace("\n", " ")
        lines.append(
            f"| {shot.get('index', '')} | {scene_name} | {shot.get('duration', '?')}s | "
            f"{narration} | {visual} | {shot.get('camera_movement', '')} | {action} | {shot.get('mood', '')} |"
        )
    lines.extend([
        "",
        "## 可灵生成流程",
        "",
        "1. 先按「角色设定」生成各角色设定卡，作为可灵主体参考",
        "2. 按「场景设定」生成各场景的视觉参考（九宫格）",
        "3. 同一场景内的镜头连续生成，复用角色卡与场景参考以保持一致性",
        "4. 把上方「完整旁白/对话」粘贴到剪辑软件做 AI 配音",
        "5. 按「镜头运动」列设置运镜，按「情绪」列匹配 BGM",
        "6. 导出 1080p，发布到 B站",
    ])
    return "\n".join(lines)


# ===================== Word 文档生成（.docx）=====================

def _get_docx():
    """延迟导入 python-docx，未安装时给出友好提示。"""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        return Document, Pt, Inches, Cm, RGBColor, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT
    except ImportError:
        raise RuntimeError("未安装 python-docx 包，请 pip install python-docx")


def _set_cell_font(cell, text, bold=False, size=10, color=None):
    """设置单元格字体（支持中文）。"""
    Document, Pt, Inches, Cm, RGBColor, _, _ = _get_docx()
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", "Microsoft YaHei")
    if color:
        run.font.color.rgb = RGBColor(*color)


def script_to_docx(
    topic: Dict[str, Any],
    script: Dict[str, Any],
    clips_map: Dict[int, List[Dict[str, Any]]] = None,
) -> "bytes":
    """
    把叙事短片剧本转成 Word 文档（.docx），返回 BytesIO。
    包含：元信息、角色设定、场景设定、完整旁白/对话、分镜表（含可灵视觉提示）、生成流程。
    """
    import io
    Document, Pt, Inches, Cm, RGBColor, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT = _get_docx()

    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", "Microsoft YaHei")

    # === 标题 ===
    title = doc.add_heading(script.get("title", topic.get("title", "未命名短片")), level=0)
    for run in title.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", "Microsoft YaHei")

    # === 元信息 ===
    doc.add_paragraph(f"赛道：{config.NICHE}")
    doc.add_paragraph(f"题材：{script.get('genre', topic.get('genre', '未知'))}")
    doc.add_paragraph(f"主题：{script.get('theme', '')}")
    doc.add_paragraph(f"对应热点：{topic.get('hot_source', '无')}")
    planned_dur = script.get("total_duration", "?")
    doc.add_paragraph(f"计划时长：{planned_dur} 秒")
    if script.get("actual_total_duration"):
        actual_dur = script.get("actual_total_duration")
        diff = script.get("duration_diff", 0)
        p = doc.add_paragraph()
        run = p.add_run(f"实际时长（根据素材调整）：{actual_dur} 秒")
        run.font.bold = True
        if diff > 0:
            run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)  # 红色：变长
        elif diff < 0:
            run.font.color.rgb = RGBColor(0x34, 0x98, 0xDB)  # 蓝色：变短
        else:
            run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)  # 绿色：一致
        p.add_run(f"（原计划 {planned_dur} 秒，差 {diff:+d} 秒）")
    doc.add_paragraph(f"BGM 情绪标签：{script.get('bgm_mood', 'mysterious')}")
    doc.add_paragraph(f"镜头数量：{len(script.get('shots', []))} 个 / 场景 {len(script.get('scenes', []))} 个 / 角色 {len(script.get('characters', []))} 个")

    doc.add_paragraph()  # 空行

    # === 角色设定 ===
    doc.add_heading("角色设定", level=1)
    for c in script.get("characters", []):
        doc.add_heading(f"{c.get('name', '未命名')}（{c.get('role', '角色')}）", level=2)
        doc.add_paragraph(f"外貌：{c.get('appearance', '')}")
        doc.add_paragraph(f"性格：{c.get('personality', '')}")
        doc.add_paragraph(f"声音：{c.get('voice', '')}")

    doc.add_paragraph()  # 空行

    # === 场景设定 ===
    doc.add_heading("场景设定", level=1)
    for s in script.get("scenes", []):
        doc.add_heading(f"场景{s.get('id', '?')}：{s.get('name', '未命名')}", level=2)
        doc.add_paragraph(f"地点：{s.get('location', '')}")
        doc.add_paragraph(f"时段：{s.get('time_of_day', '')}")
        doc.add_paragraph(f"天气/氛围：{s.get('weather', '')}")
        doc.add_paragraph(f"视觉描述：{s.get('description', '')}")

    doc.add_paragraph()  # 空行

    # === 完整旁白/对话 ===
    doc.add_heading("完整旁白/对话（可直接复制做 AI 配音）", level=1)
    narration = script.get("narration_full", "")
    p = doc.add_paragraph(narration)
    for run in p.runs:
        run.font.size = Pt(12)

    doc.add_paragraph()  # 空行

    # === 分镜表 ===
    doc.add_heading("分镜详情表（含可灵视觉提示）", level=1)
    headers = ["序号", "场景", "时长", "旁白/对话", "画面描述", "镜头运动", "角色动作", "情绪"]
    col_widths = [Cm(1.0), Cm(2.0), Cm(1.2), Cm(4.5), Cm(4.5), Cm(1.8), Cm(3.0), Cm(1.8)]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    hdr_cells = table.rows[0].cells
    for i, (header, width) in enumerate(zip(headers, col_widths)):
        hdr_cells[i].width = width
        _set_cell_font(hdr_cells[i], header, bold=True, size=10, color=(0xFF, 0xFF, 0xFF))

    # 数据行
    for shot in script.get("shots", []):
        row_cells = table.add_row().cells
        data = [
            str(shot.get("index", "")),
            shot.get("scene_name", ""),
            f"{shot.get('duration', '?')}s",
            shot.get("narration", ""),
            shot.get("visual_description", ""),
            shot.get("camera_movement", ""),
            shot.get("character_action", ""),
            shot.get("mood", ""),
        ]
        for i, (cell, text, width) in enumerate(zip(row_cells, data, col_widths)):
            cell.width = width
            # 序号、场景、时长、镜头运动、情绪列加粗居中
            if i in (0, 1, 2, 5, 7):
                _set_cell_font(cell, text, bold=True, size=9)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                _set_cell_font(cell, text, size=9)

    doc.add_paragraph()  # 空行

    # === 可灵生成流程 ===
    doc.add_heading("可灵生成流程", level=1)
    steps = [
        "1. 先按「角色设定」生成各角色设定卡，作为可灵主体参考",
        "2. 按「场景设定」生成各场景的视觉参考（九宫格）",
        "3. 同一场景内的镜头连续生成，复用角色卡与场景参考以保持一致性",
        "4. 把上方「完整旁白/对话」粘贴到剪辑软件做 AI 配音",
        "5. 按「镜头运动」列设置运镜，按「情绪」列匹配 BGM",
        "6. 导出 1080p，发布到 B站",
    ]
    for step in steps:
        doc.add_paragraph(step)

    # === 字数匹配说明 ===
    doc.add_heading("旁白/对话与时长匹配说明", level=1)
    cps = config.NARRATION_CHARS_PER_SEC
    doc.add_paragraph(f"中文语速基准：{cps} 字/秒。每个镜头的旁白/对白字数 ≈ 时长 × {cps}。")
    doc.add_paragraph("如果实际生成的镜头时长与计划不符，请按以下方案处理：")
    doc.add_paragraph("  ✅ 镜头够长：截取精华段落，剪掉冗余")
    doc.add_paragraph("  ⚠️ 镜头略短：末尾慢放或补一拍静帧")
    doc.add_paragraph("  ⚠️ 镜头太短：精简旁白字数或补空镜过渡")

    # 输出到 BytesIO
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def publish_info_to_docx(info: Dict[str, Any]) -> "bytes":
    """把 B站发布信息转成 Word 文档（.docx），返回 BytesIO。"""
    import io
    Document, Pt, Inches, Cm, RGBColor, WD_ALIGN_PARAGRAPH, _ = _get_docx()

    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", "Microsoft YaHei")

    # === 标题 ===
    heading = doc.add_heading("B站发布文案（直接复制使用）", level=0)
    for run in heading.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", "Microsoft YaHei")

    # === 作品标题 ===
    doc.add_heading("作品标题", level=1)
    p = doc.add_paragraph()
    run = p.add_run(info.get("title", ""))
    run.font.bold = True
    run.font.size = Pt(14)

    # === 作品简介 ===
    doc.add_heading("作品简介", level=1)
    doc.add_paragraph(info.get("description", ""))

    # === 标签 ===
    doc.add_heading("标签（发布时填到标签栏）", level=1)
    p = doc.add_paragraph(" ".join(info.get("hashtags", [])))
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x00, 0x77, 0xCC)

    # === 分区推荐 ===
    doc.add_heading("分区推荐", level=1)
    p = doc.add_paragraph(info.get("partition", ""))
    for run in p.runs:
        run.font.bold = True
        run.font.size = Pt(12)

    # === 封面建议 ===
    doc.add_heading("封面建议", level=1)
    doc.add_paragraph(f"封面文字突出：{info.get('cover_suggestion', '')}")

    # === 发布建议 ===
    doc.add_heading("发布建议", level=1)
    doc.add_paragraph(info.get("publish_tips", ""))

    # === 复制清单 ===
    doc.add_heading("复制清单", level=1)

    doc.add_heading("标题（复制到标题栏）", level=2)
    p = doc.add_paragraph(info.get("title", ""))
    p.style = doc.styles["Normal"]
    for run in p.runs:
        run.font.size = Pt(12)

    doc.add_heading("简介（复制到简介栏）", level=2)
    p = doc.add_paragraph(info.get("description", ""))
    for run in p.runs:
        run.font.size = Pt(12)

    doc.add_heading("标签（复制到标签栏）", level=2)
    p = doc.add_paragraph(" ".join(info.get("hashtags", [])))
    for run in p.runs:
        run.font.size = Pt(12)

    # 输出到 BytesIO
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


if __name__ == "__main__":
    # 自测：用 mock 热点跑一遍
    mock_hot = [
        {"title": "深海发现新物种发光水母", "source": "zhihu"},
        {"title": "某明星官宣结婚", "source": "weibo"},
        {"title": "火星探测新突破发现液态水痕迹", "source": "douyin"},
    ]
    print("=== 测试选题生成 ===")
    topics = generate_topics(mock_hot, count=2)
    print(json.dumps(topics, ensure_ascii=False, indent=2))

    if topics:
        print("\n=== 测试脚本生成 ===")
        script = generate_script(topics[0])
        print(json.dumps(script, ensure_ascii=False, indent=2))
        print("\n=== Markdown 输出 ===")
        print(script_to_markdown(topics[0], script))
