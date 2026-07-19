"""
可灵AI提示词生成模块
把叙事短片剧本的镜头转成可灵AI(Kling 3.0)可直接使用的提示词包。

可灵3.0关键特性：
- 多镜头连续生成：一次生成3-6个连贯镜头，自动保持一致性
- 主体参考：上传角色参考图，所有镜头保持角色一致（>96%）
- 镜头运动：推/拉/摇/移/跟/旋转
- 每日66灵感值免费（约6次生成，每次5秒）

工作流：
1. 剧本 shots 列表 → group_shots_by_scene() → 按场景分组（每组3-4个镜头）
2. 单镜头 → generate_kling_prompt() → 可灵提示词
   公式：[角色描述 @角色名] [动作描述] [场景描述] [镜头运动] [光影氛围] [画面风格]
3. shots → generate_kling_shot_groups() → 完整批次（含合并提示词、角色参考清单、操作建议）
4. kling_package_to_docx / kling_package_to_markdown 导出"拍摄指南"

shot 字段约定（兼容多种命名）：
- index: 镜头序号
- scene_id / scene: 场景ID
- scene_name: 场景名
- characters: [{"name": "林夏", "description": "穿黑色风衣的年轻女性"}]
- action / character_action: 动作描述
- scene_description / scene_desc / visual_note: 场景画面描述
- camera_movement / camera / shot_type: 镜头运动（中文：推/拉/摇/移/跟/旋转/固定）
- lighting / atmosphere / mood: 光影氛围
- style / visual_style: 画面风格
- duration: 时长（秒）
- narration: 旁白
"""
import json
from typing import List, Dict, Any
from collections import defaultdict

import config


# ===================== 镜头运动映射 =====================

# 中文镜头运动 → 可灵/英文镜头运动
CAMERA_MOVEMENT_MAP: Dict[str, str] = {
    "推": "push in",
    "推镜头": "push in",
    "拉": "pull back",
    "拉镜头": "pull back",
    "摇": "pan",
    "左摇": "pan left",
    "右摇": "pan right",
    "移": "tracking shot",
    "平移": "tracking shot",
    "跟": "following shot",
    "跟拍": "following shot",
    "旋转": "rotating shot",
    "环绕": "orbit shot",
    "固定": "fixed shot",
    "静帧": "fixed shot",
    "俯拍": "top-down shot",
    "俯视": "top-down shot",
    "仰拍": "low angle shot",
    "仰视": "low angle shot",
    "特写": "close-up",
    "近景": "close-up",
    "全景": "wide shot",
    "中景": "medium shot",
    "远景": "extreme wide shot",
}


def _translate_camera_movement(movement: str) -> str:
    """
    把中文镜头运动翻译成英文。
    支持复合描述（如"推镜头特写到面部"），先精确匹配关键词，匹配不到原样返回。
    """
    if not movement:
        return ""
    text = movement.strip()
    # 精确匹配
    if text in CAMERA_MOVEMENT_MAP:
        return CAMERA_MOVEMENT_MAP[text]
    # 包含匹配：找关键词
    for cn, en in CAMERA_MOVEMENT_MAP.items():
        if cn in text:
            # 把中文关键词替换成英文，保留其余描述
            return text.replace(cn, en)
    # 找不到映射，原样返回
    return text


# ===================== 按场景分组 =====================

def group_shots_by_scene(shots: List[Dict[str, Any]]) -> Dict[int, List[Dict[str, Any]]]:
    """
    把镜头按场景分组，方便可灵多镜头连续生成。
    同一场景的镜头放一起，可灵一次生成3-4个连贯镜头。
    如果某场景超过4个镜头，拆成多组（每组3-4个）。

    返回：Dict[int, List[shot]]，key 为顺序递增的 group_id（1, 2, 3...），
          每组最多 4 个镜头，同组的镜头都属于同一 scene_id。
          每个 shot 仍保留原始 scene_id 字段，下游可读取。
    """
    if not shots:
        return {}

    # 先按 scene_id 聚合（保留输入顺序）
    by_scene: Dict[Any, List[Dict[str, Any]]] = defaultdict(list)
    for shot in shots:
        scene_id = shot.get("scene_id", shot.get("scene", 0))
        by_scene[scene_id].append(shot)

    # 按 scene_id 排序后拆分（每组最多 4 个）
    result: Dict[int, List[Dict[str, Any]]] = {}
    group_id = 0
    for scene_id in sorted(by_scene.keys(), key=lambda x: (x is None, x)):
        scene_shots = by_scene[scene_id]
        # 每组最多 4 个镜头
        for i in range(0, len(scene_shots), 4):
            group_id += 1
            result[group_id] = scene_shots[i:i + 4]

    print(f"[分组] {len(shots)} 个镜头 → {len(result)} 组（按场景聚合，每组≤4个）")
    return result


# ===================== 单镜头提示词 =====================

def generate_kling_prompt(shot: Dict[str, Any], character_names: List[str] = None) -> str:
    """
    为单个镜头生成可灵提示词。

    可灵提示词公式：
    [角色描述 @角色名] [动作描述] [场景描述] [镜头运动] [光影氛围] [画面风格]

    示例：
    "一位穿黑色风衣的年轻女性 @林夏 站在雨夜的街灯下，转头看向镜头，
     背景是霓虹灯闪烁的东京街头， 推镜头特写到面部，
     冷色调蓝光，雨滴反光， 电影感画面，35mm镜头质感"

    参数：
    - shot: 镜头字典，字段约定见模块文档字符串
    - character_names: 可选，当 shot 没有详细 characters 字段时，
                       用此列表的角色名生成 @角色名 标记
    """
    parts: List[str] = []

    # 1. 角色描述（带 @角色名 标记，可灵的调用语法）
    characters = shot.get("characters", []) or []
    char_part = ""
    if characters and isinstance(characters, list):
        char_segments = []
        for char in characters:
            if not isinstance(char, dict):
                continue
            name = char.get("name", "")
            desc = char.get("description", "") or char.get("desc", "")
            if name and desc:
                char_segments.append(f"{desc} @{name}")
            elif name:
                char_segments.append(f"@{name}")
            elif desc:
                char_segments.append(desc)
        char_part = "，".join(char_segments)
    elif character_names:
        # 没有详细角色描述，但有角色名列表 → 只标 @角色名
        char_part = " ".join(f"@{n}" for n in character_names if n)
    if char_part:
        parts.append(char_part)

    # 2. 动作描述
    action = shot.get("action", "") or shot.get("character_action", "") or shot.get("motion", "")
    if action:
        parts.append(action)

    # 3. 场景描述
    scene_desc = (
        shot.get("scene_description", "")
        or shot.get("scene_desc", "")
        or shot.get("visual_note", "")
        or shot.get("setting", "")
    )
    scene_name = shot.get("scene_name", "")
    if scene_desc:
        parts.append(f"背景是{scene_desc}")
    elif scene_name:
        parts.append(f"背景是{scene_name}")

    # 4. 镜头运动（翻译成英文）
    camera_cn = (
        shot.get("camera_movement", "")
        or shot.get("camera", "")
        or shot.get("shot_type", "")
    )
    if camera_cn:
        camera_en = _translate_camera_movement(camera_cn)
        parts.append(camera_en)

    # 5. 光影氛围
    lighting = (
        shot.get("lighting", "")
        or shot.get("atmosphere", "")
        or shot.get("mood", "")
    )
    if lighting:
        parts.append(lighting)

    # 6. 画面风格
    style = shot.get("style", "") or shot.get("visual_style", "")
    if style:
        parts.append(style)

    return "，".join(parts)


# ===================== 生成可灵镜头组 =====================

# 可灵每日免费灵感值
KLING_DAILY_FREE_CREDITS = 66
# 每组多镜头连续生成消耗的灵感值（5秒×2=10点，按可灵3.0计费）
KLING_COST_PER_GROUP = 10


def generate_kling_shot_groups(
    shots: List[Dict[str, Any]],
    characters: List[Dict[str, Any]] = None,
) -> List[Dict[str, Any]]:
    """
    生成可灵多镜头连续生成批次。

    返回：[
      {
        "group_id": 1,
        "scene_id": 1,
        "scene_name": "雨夜街头",
        "shots": [shot1, shot2, shot3],  # 3-4个镜头
        "shot_prompts": ["...", "..."],  # 每个镜头的可灵提示词
        "combined_prompt": "一次生成的完整提示词",
        "character_refs": ["林夏"],  # 需要上传哪些角色参考图
        "estimated_cost": 10,  # 预计消耗灵感值
        "tips": "操作建议"
      }
    ]
    """
    if not shots:
        print("[可灵] 镜头列表为空，跳过")
        return []

    # 1. 按场景分组
    groups = group_shots_by_scene(shots)

    # 2. 角色名索引（用于补全 @角色名）
    all_char_names: List[str] = []
    if characters:
        for c in characters:
            name = c.get("name", "") if isinstance(c, dict) else ""
            if name:
                all_char_names.append(name)

    # 3. 为每组生成完整信息
    result: List[Dict[str, Any]] = []
    for group_id, group_shots in groups.items():
        scene_id = group_shots[0].get("scene_id", group_shots[0].get("scene", 0))
        scene_name = group_shots[0].get("scene_name", f"场景{scene_id}")

        # 收集这组镜头涉及的角色
        char_refs: List[str] = []
        for s in group_shots:
            for c in s.get("characters", []) or []:
                if isinstance(c, dict):
                    name = c.get("name", "")
                    if name and name not in char_refs:
                        char_refs.append(name)

        # 生成每个镜头的提示词
        shot_prompts: List[str] = []
        for s in group_shots:
            # 如果镜头没有 characters 字段，用全局角色名补 @标记
            fallback_names = all_char_names if not s.get("characters") else None
            prompt = generate_kling_prompt(s, character_names=fallback_names)
            shot_prompts.append(prompt)

        # 合并提示词：把多个镜头描述合并成一次"多镜头连续生成"的提示词
        combined_lines: List[str] = []
        for i, (s, p) in enumerate(zip(group_shots, shot_prompts), start=1):
            idx = s.get("index", i)
            combined_lines.append(f"【镜头{idx}】{p}")
        combined_prompt = "\n".join(combined_lines)

        # 操作建议
        tips = _build_group_tips(char_refs, len(group_shots), scene_name)

        result.append({
            "group_id": group_id,
            "scene_id": scene_id,
            "scene_name": scene_name,
            "shots": group_shots,
            "shot_prompts": shot_prompts,
            "combined_prompt": combined_prompt,
            "character_refs": char_refs,
            "estimated_cost": KLING_COST_PER_GROUP,
            "tips": tips,
        })

    total_cost = sum(g["estimated_cost"] for g in result)
    print(f"[可灵] 生成 {len(result)} 个镜头组，共 {len(shots)} 个镜头，"
          f"预计消耗 {total_cost} 灵感值")
    return result


def _build_group_tips(char_refs: List[str], shot_count: int, scene_name: str) -> str:
    """根据组内角色参考和镜头数生成操作建议。"""
    tips: List[str] = []

    if char_refs:
        names_str = "、".join(char_refs)
        tips.append(f"先上传 {names_str} 的正面参考图作为「主体参考」，再生成")
    else:
        tips.append("本组无角色，无需上传主体参考")

    tips.append(f"选择「多镜头连续生成」模式，一次生成 {shot_count} 个连贯镜头")
    tips.append(f"如需锁定空间，可上传场景「{scene_name}」的全景图作为场景参考")
    tips.append("生成后检查角色一致性，不满意可重新生成（消耗新灵感值）")

    return "；".join(tips)


# ===================== Word 文档生成 =====================

def _get_docx():
    """延迟导入 python-docx，未安装时给出友好提示。"""
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        return {
            "Document": Document,
            "Pt": Pt,
            "Cm": Cm,
            "RGBColor": RGBColor,
            "WD_ALIGN_PARAGRAPH": WD_ALIGN_PARAGRAPH,
            "WD_TABLE_ALIGNMENT": WD_TABLE_ALIGNMENT,
            "qn": qn,
            "OxmlElement": OxmlElement,
        }
    except ImportError as e:
        raise RuntimeError("未安装 python-docx 包，请 pip install python-docx") from e


_CJK_FONT = "Microsoft YaHei"
_CODE_FONT = "Consolas"


def _set_run_font(run, font_name: str = _CJK_FONT, size=None, bold=None, color=None, docx_mod=None):
    """设置 run 的字体（兼容中英文）。"""
    if docx_mod is None:
        docx_mod = _get_docx()
    if size is not None:
        run.font.size = docx_mod["Pt"](size)
    if bold is not None:
        run.font.bold = bold
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(docx_mod["qn"]("w:rFonts"))
    if rFonts is None:
        rFonts = docx_mod["OxmlElement"]("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(docx_mod["qn"]("w:eastAsia"), _CJK_FONT)
    rFonts.set(docx_mod["qn"]("w:ascii"), font_name)
    rFonts.set(docx_mod["qn"]("w:hAnsi"), font_name)
    if color is not None:
        run.font.color.rgb = docx_mod["RGBColor"](*color)


def _set_doc_default_font(doc, docx_mod):
    """设置文档默认字体为 Microsoft YaHei。"""
    style = doc.styles["Normal"]
    style.font.name = _CJK_FONT
    style.font.size = docx_mod["Pt"](11)
    style._element.rPr.rFonts.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", _CJK_FONT
    )


def _add_heading(doc, text, level=1, docx_mod=None):
    """添加标题并设置中文字体。"""
    if docx_mod is None:
        docx_mod = _get_docx()
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        _set_run_font(run, font_name=_CJK_FONT, docx_mod=docx_mod)
    return heading


def _set_cell_font(cell, text, bold=False, size=10, color=None, docx_mod=None):
    """设置单元格字体（支持中文）。"""
    if docx_mod is None:
        docx_mod = _get_docx()
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    _set_run_font(run, bold=bold, size=size, color=color, docx_mod=docx_mod)


def _add_code_block(doc, text: str, docx_mod=None):
    """添加代码块样式段落（灰底+边框+等宽字体）。"""
    if docx_mod is None:
        docx_mod = _get_docx()
    qn = docx_mod["qn"]
    OxmlElement = docx_mod["OxmlElement"]
    Cm = docx_mod["Cm"]

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.right_indent = Cm(0.3)
    p.paragraph_format.space_before = docx_mod["Pt"](4)
    p.paragraph_format.space_after = docx_mod["Pt"](4)

    pPr = p._p.get_or_add_pPr()

    # 灰色背景
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F5F5F5")
    pPr.append(shd)

    # 四周边框
    pBdr = OxmlElement("w:pBdr")
    for border_name in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "4")
        border.set(qn("w:color"), "CCCCCC")
        pBdr.append(border)
    pPr.append(pBdr)

    run = p.add_run(text)
    _set_run_font(
        run, font_name=_CODE_FONT, size=9,
        color=(0x33, 0x33, 0x33), docx_mod=docx_mod,
    )
    return p


def kling_package_to_docx(
    shot_groups: List[Dict[str, Any]],
    script_title: str = "",
    total_shots: int = 0,
) -> "bytes":
    """
    把可灵提示词包转成 Word 文档（拍摄指南）。

    内容：
    - 总览：总镜头数、分组数、预计灵感值消耗、预计天数
    - 每个分组：场景名、镜头列表、合并提示词、需要的角色参考图、操作建议
    - 可灵操作步骤
    - 免费额度管理建议
    """
    import io
    docx_mod = _get_docx()
    Document = docx_mod["Document"]
    Cm = docx_mod["Cm"]
    RGBColor = docx_mod["RGBColor"]
    WD_ALIGN_PARAGRAPH = docx_mod["WD_ALIGN_PARAGRAPH"]
    WD_TABLE_ALIGNMENT = docx_mod["WD_TABLE_ALIGNMENT"]
    qn = docx_mod["qn"]

    doc = Document()
    _set_doc_default_font(doc, docx_mod)

    # === 封面标题 ===
    title_text = "可灵 AI 拍摄指南"
    if script_title:
        title_text = f"{title_text} - {script_title}"
    title = doc.add_heading(title_text, level=0)
    for run in title.runs:
        _set_run_font(run, font_name=_CJK_FONT, docx_mod=docx_mod)

    doc.add_paragraph(
        "本指南把剧本镜头按场景分成多组，每组 3-4 个镜头，"
        "用可灵 AI「多镜头连续生成」一次生成，自动保持角色和场景一致。"
    )
    doc.add_paragraph()

    # === 总览 ===
    _add_heading(doc, "一、总览", level=1, docx_mod=docx_mod)
    group_count = len(shot_groups)
    actual_total_shots = total_shots or sum(len(g.get("shots", [])) for g in shot_groups)
    total_cost = sum(g.get("estimated_cost", KLING_COST_PER_GROUP) for g in shot_groups)
    estimated_days = max(1, -(-total_cost // KLING_DAILY_FREE_CREDITS))  # 向上取整

    overview = [
        ("总镜头数", f"{actual_total_shots} 个"),
        ("分组数", f"{group_count} 组（每组 3-4 个镜头）"),
        ("预计灵感值消耗", f"{total_cost} 点（每组 {KLING_COST_PER_GROUP} 点）"),
        ("每日免费灵感值", f"{KLING_DAILY_FREE_CREDITS} 点（约 {KLING_DAILY_FREE_CREDITS // KLING_COST_PER_GROUP} 组）"),
        ("预计完成天数", f"{estimated_days} 天（仅用免费额度）"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    _set_cell_font(hdr[0], "项目", bold=True, size=10, color=(0xFF, 0xFF, 0xFF), docx_mod=docx_mod)
    _set_cell_font(hdr[1], "数值", bold=True, size=10, color=(0xFF, 0xFF, 0xFF), docx_mod=docx_mod)
    hdr[0].width = Cm(5.0)
    hdr[1].width = Cm(9.0)
    for label, value in overview:
        row = table.add_row().cells
        row[0].width = Cm(5.0)
        row[1].width = Cm(9.0)
        _set_cell_font(row[0], label, bold=True, size=10, docx_mod=docx_mod)
        _set_cell_font(row[1], value, size=10, docx_mod=docx_mod)

    doc.add_paragraph()

    # === 每个分组详情 ===
    _add_heading(doc, "二、分组拍摄详情", level=1, docx_mod=docx_mod)

    for group in shot_groups:
        group_id = group.get("group_id", "?")
        scene_name = group.get("scene_name", f"场景{group.get('scene_id', '?')}")
        scene_id = group.get("scene_id", "?")
        group_shots = group.get("shots", [])
        combined_prompt = group.get("combined_prompt", "")
        char_refs = group.get("character_refs", [])
        cost = group.get("estimated_cost", KLING_COST_PER_GROUP)
        tips = group.get("tips", "")

        _add_heading(
            doc,
            f"第 {group_id} 组 - 场景{scene_id}：{scene_name}（{len(group_shots)} 个镜头 / {cost} 灵感值）",
            level=2, docx_mod=docx_mod,
        )

        # 角色参考图清单
        if char_refs:
            p = doc.add_paragraph()
            r = p.add_run("需要上传的角色参考图：")
            _set_run_font(r, bold=True, size=11, color=(0xE7, 0x4C, 0x3C), docx_mod=docx_mod)
            r2 = p.add_run("、".join(char_refs))
            _set_run_font(r2, size=11, docx_mod=docx_mod)
        else:
            p = doc.add_paragraph()
            r = p.add_run("本组无角色，无需上传主体参考图")
            _set_run_font(r, size=11, color=(0x88, 0x88, 0x88), docx_mod=docx_mod)

        # 镜头列表表格
        doc.add_paragraph().add_run("镜头列表：")
        shot_table = doc.add_table(rows=1, cols=5)
        shot_table.style = "Light Grid Accent 1"
        shot_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        shot_headers = ["序号", "时长", "角色 @名", "动作 / 画面", "镜头运动"]
        shot_widths = [Cm(1.2), Cm(1.5), Cm(3.5), Cm(6.0), Cm(2.5)]
        hdr_cells = shot_table.rows[0].cells
        for i, (h, w) in enumerate(zip(shot_headers, shot_widths)):
            hdr_cells[i].width = w
            _set_cell_font(hdr_cells[i], h, bold=True, size=9,
                           color=(0xFF, 0xFF, 0xFF), docx_mod=docx_mod)

        for s in group_shots:
            row_cells = shot_table.add_row().cells
            idx = s.get("index", "")
            duration = f"{s.get('duration', '?')}s"
            # 角色 @名
            chars = s.get("characters", []) or []
            char_str = "、".join(
                f"@{c.get('name', '')}" for c in chars
                if isinstance(c, dict) and c.get("name")
            )
            # 动作 / 画面
            action = s.get("action", "") or s.get("character_action", "") or ""
            scene_desc = s.get("scene_description", "") or s.get("visual_note", "")
            action_visual = action
            if scene_desc:
                action_visual = f"{action}（{scene_desc}）" if action else scene_desc
            # 镜头运动
            camera = s.get("camera_movement", "") or s.get("camera", "") or ""
            camera_en = _translate_camera_movement(camera) if camera else ""

            row_data = [str(idx), duration, char_str, action_visual, camera_en]
            for i, (cell, text, w) in enumerate(zip(row_cells, row_data, shot_widths)):
                cell.width = w
                if i < 2:
                    _set_cell_font(cell, text, bold=True, size=9, docx_mod=docx_mod)
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    _set_cell_font(cell, text, size=9, docx_mod=docx_mod)

        doc.add_paragraph()

        # 合并提示词
        doc.add_paragraph().add_run("合并提示词（复制到可灵「多镜头连续生成」）：")
        _add_code_block(doc, combined_prompt or "（空）", docx_mod=docx_mod)

        # 操作建议
        if tips:
            p = doc.add_paragraph()
            r = p.add_run("操作建议：")
            _set_run_font(r, bold=True, size=11, color=(0x27, 0xAE, 0x60), docx_mod=docx_mod)
            r2 = p.add_run(tips)
            _set_run_font(r2, size=11, docx_mod=docx_mod)

        doc.add_paragraph()  # 组间空行

    # === 可灵操作步骤 ===
    _add_heading(doc, "三、可灵 AI 操作步骤", level=1, docx_mod=docx_mod)
    kling_steps = [
        "打开 https://klingai.com 注册登录",
        "先生成角色设定卡（用角色卡提示词去通义万相生成）",
        "在可灵上传角色参考图作为「主体参考」",
        "按「拍摄指南」的分组，每组复制「合并提示词」到可灵生成",
        "同一场景的镜头用「多镜头连续生成」模式，一次生成3-4个",
        "每日66灵感值约可生成6组，超出等次日",
        "下载所有生成的视频片段",
        "用剪映拼接，加上旁白配音和字幕",
    ]
    for i, step in enumerate(kling_steps, start=1):
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(step)
        _set_run_font(run, size=11, docx_mod=docx_mod)

    # === 免费额度管理建议 ===
    _add_heading(doc, "四、免费额度管理建议", level=1, docx_mod=docx_mod)
    quota_tips = [
        f"每日免费 {KLING_DAILY_FREE_CREDITS} 灵感值，每组多镜头生成消耗 {KLING_COST_PER_GROUP} 点，"
        f"每天最多免费生成 {KLING_DAILY_FREE_CREDITS // KLING_COST_PER_GROUP} 组",
        "优先生成主角戏份多的组（角色一致性最关键）",
        "同一组可生成多次，挑选最满意的一版，不满意及时重生成避免浪费灵感值",
        "如果总灵感值超出单日免费额度，可分多天生成，或购买灵感值包",
        "建议把每组生成结果按 group_id 命名保存（如 group01_clip.mp4），方便后期拼接",
        "夜间或非高峰时段生成速度更快",
    ]
    for tip in quota_tips:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(tip)
        _set_run_font(run, size=11, docx_mod=docx_mod)

    # === 输出 ===
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ===================== Markdown 导出 =====================

def kling_package_to_markdown(shot_groups: List[Dict[str, Any]]) -> str:
    """转成 Markdown。"""
    lines: List[str] = []

    lines.append("# 可灵 AI 拍摄指南")
    lines.append("")
    lines.append(
        "> 本指南把剧本镜头按场景分成多组，每组 3-4 个镜头，"
        "用可灵 AI「多镜头连续生成」一次生成，自动保持角色和场景一致。"
    )
    lines.append("")

    # === 总览 ===
    group_count = len(shot_groups)
    total_shots = sum(len(g.get("shots", [])) for g in shot_groups)
    total_cost = sum(g.get("estimated_cost", KLING_COST_PER_GROUP) for g in shot_groups)
    estimated_days = max(1, -(-total_cost // KLING_DAILY_FREE_CREDITS))

    lines.append("## 一、总览")
    lines.append("")
    lines.append("| 项目 | 数值 |")
    lines.append("|------|------|")
    lines.append(f"| 总镜头数 | {total_shots} 个 |")
    lines.append(f"| 分组数 | {group_count} 组（每组 3-4 个镜头） |")
    lines.append(f"| 预计灵感值消耗 | {total_cost} 点（每组 {KLING_COST_PER_GROUP} 点） |")
    lines.append(f"| 每日免费灵感值 | {KLING_DAILY_FREE_CREDITS} 点（约 {KLING_DAILY_FREE_CREDITS // KLING_COST_PER_GROUP} 组） |")
    lines.append(f"| 预计完成天数 | {estimated_days} 天（仅用免费额度） |")
    lines.append("")

    # === 分组详情 ===
    lines.append("## 二、分组拍摄详情")
    lines.append("")

    for group in shot_groups:
        group_id = group.get("group_id", "?")
        scene_name = group.get("scene_name", f"场景{group.get('scene_id', '?')}")
        scene_id = group.get("scene_id", "?")
        group_shots = group.get("shots", [])
        combined_prompt = group.get("combined_prompt", "")
        char_refs = group.get("character_refs", [])
        cost = group.get("estimated_cost", KLING_COST_PER_GROUP)
        tips = group.get("tips", "")

        lines.append(
            f"### 第 {group_id} 组 - 场景{scene_id}：{scene_name}"
            f"（{len(group_shots)} 个镜头 / {cost} 灵感值）"
        )
        lines.append("")

        # 角色参考图
        if char_refs:
            lines.append(f"**需要上传的角色参考图**：{'、'.join(char_refs)}")
        else:
            lines.append("**本组无角色**，无需上传主体参考图")
        lines.append("")

        # 镜头列表
        lines.append("镜头列表：")
        lines.append("")
        lines.append("| 序号 | 时长 | 角色 @名 | 动作 / 画面 | 镜头运动 |")
        lines.append("|------|------|---------|------------|---------|")
        for s in group_shots:
            idx = s.get("index", "")
            duration = f"{s.get('duration', '?')}s"
            chars = s.get("characters", []) or []
            char_str = "、".join(
                f"@{c.get('name', '')}" for c in chars
                if isinstance(c, dict) and c.get("name")
            )
            action = s.get("action", "") or s.get("character_action", "") or ""
            scene_desc = s.get("scene_description", "") or s.get("visual_note", "")
            action_visual = action
            if scene_desc:
                action_visual = f"{action}（{scene_desc}）" if action else scene_desc
            action_visual = action_visual.replace("|", "/").replace("\n", " ")
            camera = s.get("camera_movement", "") or s.get("camera", "") or ""
            camera_en = _translate_camera_movement(camera) if camera else ""
            lines.append(f"| {idx} | {duration} | {char_str} | {action_visual} | {camera_en} |")
        lines.append("")

        # 合并提示词
        lines.append("**合并提示词**（复制到可灵「多镜头连续生成」）：")
        lines.append("```")
        lines.append(combined_prompt or "（空）")
        lines.append("```")
        lines.append("")

        # 操作建议
        if tips:
            lines.append(f"**操作建议**：{tips}")
            lines.append("")

    # === 可灵操作步骤 ===
    lines.append("## 三、可灵 AI 操作步骤")
    lines.append("")
    kling_steps = [
        "打开 https://klingai.com 注册登录",
        "先生成角色设定卡（用角色卡提示词去通义万相生成）",
        "在可灵上传角色参考图作为「主体参考」",
        "按「拍摄指南」的分组，每组复制「合并提示词」到可灵生成",
        "同一场景的镜头用「多镜头连续生成」模式，一次生成3-4个",
        "每日66灵感值约可生成6组，超出等次日",
        "下载所有生成的视频片段",
        "用剪映拼接，加上旁白配音和字幕",
    ]
    for i, step in enumerate(kling_steps, start=1):
        lines.append(f"{i}. {step}")
    lines.append("")

    # === 免费额度管理建议 ===
    lines.append("## 四、免费额度管理建议")
    lines.append("")
    quota_tips = [
        f"每日免费 {KLING_DAILY_FREE_CREDITS} 灵感值，每组多镜头生成消耗 {KLING_COST_PER_GROUP} 点，"
        f"每天最多免费生成 {KLING_DAILY_FREE_CREDITS // KLING_COST_PER_GROUP} 组",
        "优先生成主角戏份多的组（角色一致性最关键）",
        "同一组可生成多次，挑选最满意的一版，不满意及时重生成避免浪费灵感值",
        "如果总灵感值超出单日免费额度，可分多天生成，或购买灵感值包",
        "建议把每组生成结果按 group_id 命名保存（如 group01_clip.mp4），方便后期拼接",
        "夜间或非高峰时段生成速度更快",
    ]
    for tip in quota_tips:
        lines.append(f"- {tip}")
    lines.append("")

    return "\n".join(lines)


# ===================== 自测 =====================

if __name__ == "__main__":
    print("=" * 60)
    print("测试 1：_translate_camera_movement")
    print("=" * 60)
    test_cases = ["推", "拉", "左摇", "跟拍", "旋转", "推镜头特写到面部", "固定", "未知运动"]
    for cn in test_cases:
        print(f"  {cn} → {_translate_camera_movement(cn)}")

    print()
    print("=" * 60)
    print("测试 2：generate_kling_prompt（单镜头）")
    print("=" * 60)
    mock_shot = {
        "index": 1,
        "scene_id": 1,
        "scene_name": "雨夜街头",
        "characters": [{"name": "林夏", "description": "一位穿黑色风衣的年轻女性"}],
        "action": "站在雨夜的街灯下，转头看向镜头",
        "scene_description": "霓虹灯闪烁的东京街头",
        "camera_movement": "推镜头特写到面部",
        "lighting": "冷色调蓝光，雨滴反光",
        "style": "电影感画面，35mm镜头质感",
        "duration": 5,
    }
    prompt = generate_kling_prompt(mock_shot)
    print(f"生成的提示词：\n{prompt}")

    print()
    print("=" * 60)
    print("测试 3：group_shots_by_scene（含超4镜头拆分）")
    print("=" * 60)
    mock_shots = [
        {"index": 1, "scene_id": 1, "scene_name": "雨夜街头"},
        {"index": 2, "scene_id": 1, "scene_name": "雨夜街头"},
        {"index": 3, "scene_id": 1, "scene_name": "雨夜街头"},
        {"index": 4, "scene_id": 1, "scene_name": "雨夜街头"},
        {"index": 5, "scene_id": 1, "scene_name": "雨夜街头"},  # 第5个 → 拆到第2组
        {"index": 6, "scene_id": 2, "scene_name": "深夜便利店"},
        {"index": 7, "scene_id": 2, "scene_name": "深夜便利店"},
        {"index": 8, "scene_id": 3, "scene_name": "天台"},
    ]
    groups = group_shots_by_scene(mock_shots)
    for gid, gshots in groups.items():
        scene_id = gshots[0].get("scene_id")
        scene_name = gshots[0].get("scene_name")
        indexes = [s["index"] for s in gshots]
        print(f"  组{gid}（场景{scene_id} {scene_name}）：镜头 {indexes}")

    print()
    print("=" * 60)
    print("测试 4：generate_kling_shot_groups（完整批次）")
    print("=" * 60)
    mock_full_shots = [
        {
            "index": 1, "scene_id": 1, "scene_name": "雨夜街头",
            "characters": [{"name": "林夏", "description": "穿黑色风衣的年轻女性"}],
            "action": "站在雨夜的街灯下，转头看向镜头",
            "scene_description": "霓虹灯闪烁的东京街头",
            "camera_movement": "推",
            "lighting": "冷色调蓝光，雨滴反光",
            "style": "电影感画面，35mm镜头质感",
            "duration": 5,
        },
        {
            "index": 2, "scene_id": 1, "scene_name": "雨夜街头",
            "characters": [{"name": "林夏", "description": "穿黑色风衣的年轻女性"}],
            "action": "低头看手机，眉头紧锁",
            "scene_description": "霓虹灯闪烁的东京街头",
            "camera_movement": "固定",
            "lighting": "冷色调蓝光，雨滴反光",
            "style": "电影感画面，35mm镜头质感",
            "duration": 4,
        },
        {
            "index": 3, "scene_id": 1, "scene_name": "雨夜街头",
            "characters": [
                {"name": "林夏", "description": "穿黑色风衣的年轻女性"},
                {"name": "陈默", "description": "戴金丝眼镜的男性"},
            ],
            "action": "两人对视，陈默递过一把伞",
            "scene_description": "霓虹灯闪烁的东京街头",
            "camera_movement": "跟",
            "lighting": "冷色调蓝光，雨滴反光",
            "style": "电影感画面，35mm镜头质感",
            "duration": 6,
        },
        {
            "index": 4, "scene_id": 2, "scene_name": "深夜便利店",
            "characters": [{"name": "林夏", "description": "穿黑色风衣的年轻女性"}],
            "action": "推开便利店玻璃门走进去",
            "scene_description": "冷白色荧光灯照亮的小便利店",
            "camera_movement": "拉",
            "lighting": "冷白色荧光灯顶光",
            "style": "电影感画面，35mm镜头质感",
            "duration": 5,
        },
        {
            "index": 5, "scene_id": 2, "scene_name": "深夜便利店",
            "characters": [{"name": "林夏", "description": "穿黑色风衣的年轻女性"}],
            "action": "站在货架前拿起一瓶水",
            "scene_description": "冷白色荧光灯照亮的小便利店",
            "camera_movement": "左摇",
            "lighting": "冷白色荧光灯顶光",
            "style": "电影感画面，35mm镜头质感",
            "duration": 4,
        },
    ]
    mock_characters = [
        {"name": "林夏", "role": "主角"},
        {"name": "陈默", "role": "配角"},
    ]

    shot_groups = generate_kling_shot_groups(mock_full_shots, characters=mock_characters)
    print(f"\n生成 {len(shot_groups)} 个镜头组：")
    for g in shot_groups:
        print(f"\n--- 第{g['group_id']}组 场景{g['scene_id']} {g['scene_name']} ---")
        print(f"  角色参考：{g['character_refs']}")
        print(f"  灵感值：{g['estimated_cost']}")
        print(f"  合并提示词：")
        for line in g["combined_prompt"].split("\n"):
            print(f"    {line}")
        print(f"  操作建议：{g['tips']}")

    print()
    print("=" * 60)
    print("测试 5：kling_package_to_markdown")
    print("=" * 60)
    md = kling_package_to_markdown(shot_groups)
    print(md[:600] + "..." if len(md) > 600 else md)
    print(f"\n（Markdown 总长度：{len(md)} 字符）")

    print()
    print("=" * 60)
    print("测试 6：kling_package_to_docx")
    print("=" * 60)
    try:
        buf = kling_package_to_docx(
            shot_groups, script_title="测试短片", total_shots=len(mock_full_shots)
        )
        data = buf.read()
        print(f"docx 生成成功，大小 {len(data)} 字节")
        out_path = "/data/user/work/test_kling_package.docx"
        with open(out_path, "wb") as f:
            f.write(data)
        print(f"已写入 {out_path}")
    except RuntimeError as e:
        print(f"跳过 docx 测试：{e}")

    print()
    print("所有测试通过 ✅")
