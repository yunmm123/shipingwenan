"""
角色和场景设定卡生成模块
为可灵AI视频生成提供角色参考图和场景参考图的绘画提示词。

工作流：
1. 剧本中的 characters 列表 → generate_character_cards() → 每个角色一张设定卡
   - card_prompt: 三视图（正/侧/背）绘画提示词，拿去通义万相/Midjourney 生成图
   - portrait_prompt: 单人正面半身像，作为可灵"主体参考"上传
   - expression_prompts: 多表情参考
2. 剧本中的 scenes 列表 → generate_scene_cards() → 每个场景一张九宫格卡
   - grid_prompt: 3x3 九宫格（9视角），锁死空间结构
   - wide_shot_prompt: 全景图，作为可灵"场景参考"上传
3. cards_to_docx / cards_to_markdown 导出可读文档

画风 art_style 影响 prompt 后缀：
- cinematic: 电影感光影、35mm镜头质感
- anime: 日系动画画风、赛璐璐
- realistic: 超写实摄影、8K细节
- oil_painting: 油画质感、厚涂笔触
- watercolor: 水彩画风格、晕染效果
"""
import json
from typing import List, Dict, Any, Union

import config

# 复用 script_generator 的 LLM 调用与 JSON 解析器
from script_generator import _chat, _robust_json_loads


# ===================== 画风映射表 =====================

ART_STYLE_PROMPTS: Dict[str, str] = {
    "cinematic": (
        "电影感光影，35mm镜头质感，写实皮肤纹理，自然景深，体积光，"
        "电影级调色，超写实细节，anamorphic lens flare"
    ),
    "anime": (
        "日系动画画风，赛璐璐上色，锐利线条，扁平光影，"
        "新海诚风格天空，宫崎骏式人物比例，anime key visual"
    ),
    "realistic": (
        "超写实摄影，8K超高细节，真实皮肤毛孔，自然光，"
        "单反相机拍摄，人像摄影级景深，photorealistic"
    ),
    "oil_painting": (
        "油画质感，厚涂笔触，古典油画色彩，伦勃朗式光影，"
        "画布纹理，oil painting texture"
    ),
    "watercolor": (
        "水彩画风格，晕染效果，淡彩通透，湿画法，"
        "纸张纹理，柔和水彩边缘，watercolor illustration"
    ),
}

# 默认画风兜底
DEFAULT_ART_STYLE = "cinematic"


def _style_suffix(art_style: str) -> str:
    """取画风后缀提示词，未知画风回退到 cinematic。"""
    return ART_STYLE_PROMPTS.get(art_style, ART_STYLE_PROMPTS[DEFAULT_ART_STYLE])


# ===================== 角色设定卡 =====================

def generate_character_cards(
    characters: List[Dict[str, Any]],
    art_style: str = "cinematic",
) -> List[Dict[str, Any]]:
    """
    为剧本中的每个角色生成设定卡绘画提示词。

    参数：
    - characters: 来自剧本的 characters 列表，每个元素建议含 name/role/description 等字段
    - art_style: 画风（cinematic/anime/realistic/oil_painting/watercolor）

    返回：每个角色一张设定卡，包含：
    - name: 角色名
    - role: 主角/配角
    - card_prompt: 角色设定卡完整绘画提示词（用于生成正面+侧面+背面立绘）
    - portrait_prompt: 单人正面半身像提示词（作为可灵主体参考图用）
    - expression_prompts: 3-4 个不同表情的提示词（正面/微笑/严肃/惊讶）
    - color_palette: 角色主色调建议

    设定卡要求：
    - 正面+侧面+背面三视图
    - 白色背景，无环境干扰
    - 清晰展示服装、发型、特征
    - 适合作为可灵AI"主体参考"上传
    """
    if not characters:
        print("[角色卡] 角色列表为空，跳过")
        return []

    style_hint = _style_suffix(art_style)

    system = f"""你是一位专业的角色概念设计师，擅长为AI视频生成项目制作角色设定卡。
你的任务：为每个角色生成可直接粘贴到"通义万相 / Midjourney"使用的绘画提示词。

【统一画风（必须追加到每条 prompt 末尾）】
{style_hint}

【设定卡（card_prompt）硬性要求】
1. 三视图设定卡：正面 + 侧面 + 背面 三个角度并排
2. 纯白色背景，无环境、无阴影干扰
3. 清晰展示：服装款式、发型、面部特征、配饰、鞋子
4. 全身立绘，人物居中，比例准确
5. 适合作为可灵AI"主体参考"上传（角色一致性 > 96%）
6. 提示词为一段连贯的描述，不要分点

【正面半身像（portrait_prompt）要求】
1. 单人正面半身像，胸部以上
2. 纯白色背景
3. 自然光，面部清晰可辨
4. 作为可灵"主体参考"首选图

【表情提示词（expression_prompts）要求】
生成 4 种表情：正面平静 / 微笑 / 严肃 / 惊讶
每条都是完整的单人正面半身像 prompt，只改变表情和微表情

【输出格式】严格输出 JSON 数组，每个角色一个对象：
[
  {{
    "name": "角色名",
    "role": "主角|配角",
    "card_prompt": "三视图设定卡完整提示词（一段话）",
    "portrait_prompt": "正面半身像提示词",
    "expression_prompts": {{
      "正面平静": "...",
      "微笑": "...",
      "严肃": "...",
      "惊讶": "..."
    }},
    "color_palette": "主色调建议，如：深蓝+银灰+暗红"
  }}
]

只输出 JSON，不要任何解释文字、不要 markdown 代码块标记。"""

    user = f"""角色列表（共 {len(characters)} 个）：
{json.dumps(characters, ensure_ascii=False, indent=2)}

请为每个角色生成设定卡提示词。
画风：{art_style}
画风后缀（追加到每条 prompt 末尾）：{style_hint}"""

    raw = _chat(system, user, temperature=0.7)
    cards = _robust_json_loads(raw, expect_type=list)

    # 字段补全与对齐：以输入 characters 顺序为准
    normalized: List[Dict[str, Any]] = []
    for i, src in enumerate(characters):
        card = cards[i] if i < len(cards) and isinstance(cards[i], dict) else {}
        name = card.get("name") or src.get("name") or f"角色{i + 1}"
        role = card.get("role") or src.get("role") or "配角"
        exp = card.get("expression_prompts") or {}
        if not isinstance(exp, dict):
            exp = {}
        # 保证 4 种表情键存在
        for key in ["正面平静", "微笑", "严肃", "惊讶"]:
            exp.setdefault(key, "")
        normalized.append({
            "name": name,
            "role": role,
            "card_prompt": card.get("card_prompt", ""),
            "portrait_prompt": card.get("portrait_prompt", ""),
            "expression_prompts": exp,
            "color_palette": card.get("color_palette", ""),
        })

    print(f"[角色卡] 生成 {len(normalized)} 张角色设定卡，画风={art_style}")
    for c in normalized:
        print(f"  - {c['name']}（{c['role']}）主色：{c['color_palette'] or '未指定'}")
    return normalized


# ===================== 场景九宫格设定卡 =====================

def generate_scene_cards(
    scenes: List[Dict[str, Any]],
    art_style: str = "cinematic",
) -> List[Dict[str, Any]]:
    """
    为剧本中的每个场景生成九宫格设定卡绘画提示词。

    参数：
    - scenes: 来自剧本的 scenes 列表，每个元素建议含 scene_id/scene_name/description 等字段
    - art_style: 画风

    返回：每个场景一张九宫格卡，包含：
    - scene_id: 场景ID
    - scene_name: 场景名称
    - grid_prompt: 九宫格绘画提示词（9个不同角度的同一空间）
    - wide_shot_prompt: 全景图提示词（作为可灵场景参考用）
    - lighting_prompt: 光影描述
    - color_palette: 场景主色调

    九宫格要求：
    - 3x3 网格，9个视角（正面/左45/右45/左90/右90/背面/俯视/仰视/特写）
    - 同一空间的9个角度，锁死空间结构
    - 统一光影和色调
    """
    if not scenes:
        print("[场景卡] 场景列表为空，跳过")
        return []

    style_hint = _style_suffix(art_style)

    system = f"""你是一位专业的场景概念设计师，擅长为AI视频生成项目制作场景设定卡。
你的任务：为每个场景生成可直接粘贴到"通义万相 / Midjourney"使用的九宫格场景设定卡提示词。

【统一画风（必须追加到每条 prompt 末尾）】
{style_hint}

【九宫格（grid_prompt）硬性要求】
1. 3x3 网格布局，共 9 个视角的同一空间：
   - 正面 / 左45° / 右45° / 左90° / 右90° / 背面 / 俯视 / 仰视 / 特写
2. 9 个视角必须是"同一个空间"的不同角度，锁死空间结构
3. 统一光影、统一色调、统一材质
4. 画面中无人物（纯场景）
5. 适合作为可灵AI"场景参考"上传，保证后续多镜头空间一致

【全景图（wide_shot_prompt）要求】
1. 单张广角全景，展现空间全貌
2. 与九宫格同一空间、同一光影
3. 作为可灵"场景参考"首选图

【输出格式】严格输出 JSON 数组，每个场景一个对象：
[
  {{
    "scene_id": 1,
    "scene_name": "场景名称",
    "grid_prompt": "九宫格场景设定卡完整提示词（一段话，含9视角描述）",
    "wide_shot_prompt": "全景图提示词",
    "lighting_prompt": "光影描述（如：冷蓝调主光+霓虹点缀，侧逆光）",
    "color_palette": "场景主色调（如：深蓝+紫红+霓虹粉）"
  }}
]

只输出 JSON，不要任何解释文字、不要 markdown 代码块标记。"""

    user = f"""场景列表（共 {len(scenes)} 个）：
{json.dumps(scenes, ensure_ascii=False, indent=2)}

请为每个场景生成九宫格场景设定卡提示词。
画风：{art_style}
画风后缀（追加到每条 prompt 末尾）：{style_hint}"""

    raw = _chat(system, user, temperature=0.7)
    cards = _robust_json_loads(raw, expect_type=list)

    # 字段补全与对齐
    normalized: List[Dict[str, Any]] = []
    for i, src in enumerate(scenes):
        card = cards[i] if i < len(cards) and isinstance(cards[i], dict) else {}
        scene_id = card.get("scene_id") or src.get("scene_id") or (i + 1)
        scene_name = card.get("scene_name") or src.get("scene_name") or f"场景{scene_id}"
        normalized.append({
            "scene_id": scene_id,
            "scene_name": scene_name,
            "grid_prompt": card.get("grid_prompt", ""),
            "wide_shot_prompt": card.get("wide_shot_prompt", ""),
            "lighting_prompt": card.get("lighting_prompt", ""),
            "color_palette": card.get("color_palette", ""),
        })

    print(f"[场景卡] 生成 {len(normalized)} 张场景设定卡，画风={art_style}")
    for s in normalized:
        print(f"  - 场景{s['scene_id']} {s['scene_name']} 主色：{s['color_palette'] or '未指定'}")
    return normalized


# ===================== Word 文档生成 =====================

def _get_docx():
    """延迟导入 python-docx，未安装时给出友好提示。"""
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        return {
            "Document": Document,
            "Pt": Pt,
            "Cm": Cm,
            "RGBColor": RGBColor,
            "WD_ALIGN_PARAGRAPH": WD_ALIGN_PARAGRAPH,
            "WD_TABLE_ALIGNMENT": WD_TABLE_ALIGNMENT,
            "qn": qn,
            "OxmlElement": OxmlElement,
        }
    except ImportError as e:
        raise RuntimeError("未安装 python-docx 包，请 pip install python-docx") from e


_CJK_FONT = "Microsoft YaHei"
_CODE_FONT = "Consolas"


def _set_run_font(run, font_name: str = _CJK_FONT, size=None, bold=None, color=None, docx_mod=None):
    """设置 run 的字体（兼容中英文）。"""
    if docx_mod is None:
        docx_mod = _get_docx()
    if size is not None:
        run.font.size = docx_mod["Pt"](size)
    if bold is not None:
        run.font.bold = bold
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(docx_mod["qn"]("w:rFonts"))
    if rFonts is None:
        rFonts = docx_mod["OxmlElement"]("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(docx_mod["qn"]("w:eastAsia"), _CJK_FONT)
    rFonts.set(docx_mod["qn"]("w:ascii"), font_name)
    rFonts.set(docx_mod["qn"]("w:hAnsi"), font_name)
    if color is not None:
        run.font.color.rgb = docx_mod["RGBColor"](*color)


def _set_doc_default_font(doc, docx_mod):
    """设置文档默认字体为 Microsoft YaHei。"""
    style = doc.styles["Normal"]
    style.font.name = _CJK_FONT
    style.font.size = docx_mod["Pt"](11)
    style._element.rPr.rFonts.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", _CJK_FONT
    )


def _add_heading(doc, text, level=1, docx_mod=None):
    """添加标题并设置中文字体。"""
    if docx_mod is None:
        docx_mod = _get_docx()
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        _set_run_font(run, font_name=_CJK_FONT, docx_mod=docx_mod)
    return heading


def _add_code_block(doc, text: str, docx_mod=None):
    """添加代码块样式段落（灰底+边框+等宽字体）。"""
    if docx_mod is None:
        docx_mod = _get_docx()
    qn = docx_mod["qn"]
    OxmlElement = docx_mod["OxmlElement"]
    Cm = docx_mod["Cm"]

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.right_indent = Cm(0.3)
    p.paragraph_format.space_before = docx_mod["Pt"](4)
    p.paragraph_format.space_after = docx_mod["Pt"](4)

    pPr = p._p.get_or_add_pPr()

    # 灰色背景
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F5F5F5")
    pPr.append(shd)

    # 四周边框
    pBdr = OxmlElement("w:pBdr")
    for border_name in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "4")
        border.set(qn("w:color"), "CCCCCC")
        pBdr.append(border)
    pPr.append(pBdr)

    run = p.add_run(text)
    _set_run_font(
        run, font_name=_CODE_FONT, size=9,
        color=(0x33, 0x33, 0x33), docx_mod=docx_mod,
    )
    return p


def cards_to_docx(
    character_cards: List,
    scene_cards: List,
    script_title: str = "",
) -> "bytes":
    """
    把角色卡和场景卡转成 Word 文档。
    包含：角色设定卡提示词、场景九宫格提示词、使用说明。
    返回 BytesIO（可直接写入文件或供下载）。
    """
    import io
    docx_mod = _get_docx()
    Document = docx_mod["Document"]
    Cm = docx_mod["Cm"]
    RGBColor = docx_mod["RGBColor"]
    WD_ALIGN_PARAGRAPH = docx_mod["WD_ALIGN_PARAGRAPH"]

    doc = Document()
    _set_doc_default_font(doc, docx_mod)

    # === 封面标题 ===
    title_text = f"角色 & 场景设定卡提示词"
    if script_title:
        title_text = f"{title_text} - {script_title}"
    title = doc.add_heading(title_text, level=0)
    for run in title.runs:
        _set_run_font(run, font_name=_CJK_FONT, docx_mod=docx_mod)

    doc.add_paragraph(
        f"用途：把下方提示词复制到「通义万相 / Midjourney」生成图，"
        f"再把生成图上传到「可灵 AI」作为主体参考 / 场景参考。"
    )
    doc.add_paragraph()

    # === 角色设定卡 ===
    _add_heading(doc, "一、角色设定卡", level=1, docx_mod=docx_mod)
    if not character_cards:
        doc.add_paragraph("（无角色）")
    else:
        doc.add_paragraph(
            f"共 {len(character_cards)} 个角色。每个角色含：三视图设定卡、正面半身像、4 种表情参考。"
        )
        doc.add_paragraph()

        for idx, card in enumerate(character_cards, start=1):
            name = card.get("name", f"角色{idx}")
            role = card.get("role", "配角")
            palette = card.get("color_palette", "未指定")

            _add_heading(doc, f"{idx}. {name}（{role}）", level=2, docx_mod=docx_mod)

            # 主色调
            p = doc.add_paragraph()
            r = p.add_run("主色调：")
            _set_run_font(r, bold=True, size=11, color=(0x55, 0x55, 0x55), docx_mod=docx_mod)
            r2 = p.add_run(palette)
            _set_run_font(r2, size=11, docx_mod=docx_mod)

            # 三视图设定卡
            doc.add_paragraph().add_run("【三视图设定卡 prompt】（生成正面+侧面+背面，白背景）")
            _add_code_block(doc, card.get("card_prompt", "") or "（空）", docx_mod=docx_mod)

            # 正面半身像
            doc.add_paragraph().add_run("【正面半身像 prompt】（可灵主体参考首选图）")
            _add_code_block(doc, card.get("portrait_prompt", "") or "（空）", docx_mod=docx_mod)

            # 表情参考
            doc.add_paragraph().add_run("【表情参考 prompt】（4 种表情，各生成一张）")
            exp = card.get("expression_prompts", {}) or {}
            for exp_name in ["正面平静", "微笑", "严肃", "惊讶"]:
                exp_text = exp.get(exp_name, "") if isinstance(exp, dict) else ""
                doc.add_paragraph().add_run(f"  · {exp_name}：")
                _add_code_block(doc, exp_text or "（空）", docx_mod=docx_mod)

            doc.add_paragraph()  # 角色间空行

    # === 场景九宫格设定卡 ===
    _add_heading(doc, "二、场景九宫格设定卡", level=1, docx_mod=docx_mod)
    if not scene_cards:
        doc.add_paragraph("（无场景）")
    else:
        doc.add_paragraph(
            f"共 {len(scene_cards)} 个场景。每个场景含：3x3 九宫格（9 视角）、全景图、光影描述。"
        )
        doc.add_paragraph()

        for idx, card in enumerate(scene_cards, start=1):
            scene_id = card.get("scene_id", idx)
            scene_name = card.get("scene_name", f"场景{scene_id}")
            palette = card.get("color_palette", "未指定")

            _add_heading(doc, f"{idx}. 场景{scene_id} - {scene_name}", level=2, docx_mod=docx_mod)

            p = doc.add_paragraph()
            r = p.add_run("主色调：")
            _set_run_font(r, bold=True, size=11, color=(0x55, 0x55, 0x55), docx_mod=docx_mod)
            r2 = p.add_run(palette)
            _set_run_font(r2, size=11, docx_mod=docx_mod)

            lighting = card.get("lighting_prompt", "")
            if lighting:
                p2 = doc.add_paragraph()
                r3 = p2.add_run("光影：")
                _set_run_font(r3, bold=True, size=11, color=(0x55, 0x55, 0x55), docx_mod=docx_mod)
                r4 = p2.add_run(lighting)
                _set_run_font(r4, size=11, docx_mod=docx_mod)

            # 九宫格
            doc.add_paragraph().add_run("【九宫格 prompt】（3x3，9 视角同一空间，锁死结构）")
            _add_code_block(doc, card.get("grid_prompt", "") or "（空）", docx_mod=docx_mod)

            # 全景图
            doc.add_paragraph().add_run("【全景图 prompt】（可灵场景参考首选图）")
            _add_code_block(doc, card.get("wide_shot_prompt", "") or "（空）", docx_mod=docx_mod)

            doc.add_paragraph()

    # === 使用说明 ===
    _add_heading(doc, "三、使用说明", level=1, docx_mod=docx_mod)

    _add_heading(doc, "1. 生成角色 / 场景参考图", level=2, docx_mod=docx_mod)
    steps_a = [
        "打开通义万相 https://tongyi.aliyun.com/wanxiang 或 Midjourney",
        "复制上方「三视图设定卡 prompt」→ 粘贴 → 生成 → 挑最准的一张存下来",
        "再复制「正面半身像 prompt」→ 生成 → 作为可灵主体参考首选图",
        "场景同理：复制「九宫格 prompt」和「全景图 prompt」分别生成",
        "建议每个角色 / 场景至少生成 4 张，挑一致性最好的",
    ]
    for s in steps_a:
        doc.add_paragraph(s, style="List Number")

    _add_heading(doc, "2. 上传到可灵 AI 作为参考", level=2, docx_mod=docx_mod)
    steps_b = [
        "打开可灵 AI https://klingai.com 并登录",
        "把角色「正面半身像」上传为「主体参考」（最多可上传 1-4 张主体参考）",
        "把场景「全景图」上传为「场景参考」（可选，用于锁定空间）",
        "后续所有该角色 / 场景的镜头生成，可灵会自动保持一致（角色一致性 > 96%）",
    ]
    for s in steps_b:
        doc.add_paragraph(s, style="List Number")

    _add_heading(doc, "3. 注意事项", level=2, docx_mod=docx_mod)
    notes = [
        "三视图设定卡必须白色背景，否则可灵主体参考会受环境干扰",
        "九宫格 9 个视角必须同一空间，否则空间结构会漂移",
        "表情参考图建议同角色同服装，只改表情",
        "生成的参考图建议 1024x1024 或以上，清晰度越高一致性越好",
    ]
    for s in notes:
        doc.add_paragraph(s, style="List Bullet")

    # === 输出 ===
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ===================== Markdown 导出 =====================

def cards_to_markdown(character_cards: List, scene_cards: List) -> str:
    """转成 Markdown 格式。"""
    lines: List[str] = []

    lines.append("# 角色 & 场景设定卡提示词")
    lines.append("")
    lines.append(
        "> 用途：把下方提示词复制到「通义万相 / Midjourney」生成图，"
        "再把生成图上传到「可灵 AI」作为主体参考 / 场景参考。"
    )
    lines.append("")

    # === 角色设定卡 ===
    lines.append("## 一、角色设定卡")
    lines.append("")
    if not character_cards:
        lines.append("（无角色）")
        lines.append("")
    else:
        lines.append(f"共 **{len(character_cards)}** 个角色。")
        lines.append("")
        for idx, card in enumerate(character_cards, start=1):
            name = card.get("name", f"角色{idx}")
            role = card.get("role", "配角")
            palette = card.get("color_palette", "未指定")

            lines.append(f"### {idx}. {name}（{role}）")
            lines.append("")
            lines.append(f"- **主色调**：{palette}")
            lines.append("")
            lines.append("**【三视图设定卡 prompt】**（生成正面+侧面+背面，白背景）")
            lines.append("```")
            lines.append(card.get("card_prompt", "") or "（空）")
            lines.append("```")
            lines.append("")
            lines.append("**【正面半身像 prompt】**（可灵主体参考首选图）")
            lines.append("```")
            lines.append(card.get("portrait_prompt", "") or "（空）")
            lines.append("```")
            lines.append("")
            lines.append("**【表情参考 prompt】**（4 种表情，各生成一张）")
            lines.append("")
            exp = card.get("expression_prompts", {}) or {}
            for exp_name in ["正面平静", "微笑", "严肃", "惊讶"]:
                exp_text = exp.get(exp_name, "") if isinstance(exp, dict) else ""
                lines.append(f"- *{exp_name}*：")
                lines.append("```")
                lines.append(exp_text or "（空）")
                lines.append("```")
                lines.append("")

    # === 场景九宫格设定卡 ===
    lines.append("## 二、场景九宫格设定卡")
    lines.append("")
    if not scene_cards:
        lines.append("（无场景）")
        lines.append("")
    else:
        lines.append(f"共 **{len(scene_cards)}** 个场景。")
        lines.append("")
        for idx, card in enumerate(scene_cards, start=1):
            scene_id = card.get("scene_id", idx)
            scene_name = card.get("scene_name", f"场景{scene_id}")
            palette = card.get("color_palette", "未指定")
            lighting = card.get("lighting_prompt", "")

            lines.append(f"### {idx}. 场景{scene_id} - {scene_name}")
            lines.append("")
            lines.append(f"- **主色调**：{palette}")
            if lighting:
                lines.append(f"- **光影**：{lighting}")
            lines.append("")
            lines.append("**【九宫格 prompt】**（3x3，9 视角同一空间，锁死结构）")
            lines.append("```")
            lines.append(card.get("grid_prompt", "") or "（空）")
            lines.append("```")
            lines.append("")
            lines.append("**【全景图 prompt】**（可灵场景参考首选图）")
            lines.append("```")
            lines.append(card.get("wide_shot_prompt", "") or "（空）")
            lines.append("```")
            lines.append("")

    # === 使用说明 ===
    lines.append("## 三、使用说明")
    lines.append("")
    lines.append("### 1. 生成角色 / 场景参考图")
    lines.append("")
    lines.append("1. 打开通义万相 https://tongyi.aliyun.com/wanxiang 或 Midjourney")
    lines.append("2. 复制上方「三视图设定卡 prompt」→ 粘贴 → 生成 → 挑最准的一张存下来")
    lines.append("3. 再复制「正面半身像 prompt」→ 生成 → 作为可灵主体参考首选图")
    lines.append("4. 场景同理：复制「九宫格 prompt」和「全景图 prompt」分别生成")
    lines.append("5. 建议每个角色 / 场景至少生成 4 张，挑一致性最好的")
    lines.append("")
    lines.append("### 2. 上传到可灵 AI 作为参考")
    lines.append("")
    lines.append("1. 打开可灵 AI https://klingai.com 并登录")
    lines.append("2. 把角色「正面半身像」上传为「主体参考」（最多 1-4 张）")
    lines.append("3. 把场景「全景图」上传为「场景参考」（可选，用于锁定空间）")
    lines.append("4. 后续所有该角色 / 场景的镜头生成，可灵会自动保持一致（角色一致性 > 96%）")
    lines.append("")
    lines.append("### 3. 注意事项")
    lines.append("")
    lines.append("- 三视图设定卡必须白色背景，否则可灵主体参考会受环境干扰")
    lines.append("- 九宫格 9 个视角必须同一空间，否则空间结构会漂移")
    lines.append("- 表情参考图建议同角色同服装，只改表情")
    lines.append("- 生成的参考图建议 1024x1024 或以上，清晰度越高一致性越好")
    lines.append("")

    return "\n".join(lines)


# ===================== 自测 =====================

if __name__ == "__main__":
    # 用 mock 数据测试核心函数（不调用 LLM）
    import sys
    from unittest.mock import patch

    # 作为 __main__ 直接运行时，_chat 绑定在 __main__ 模块上；
    # 被 import 时，_chat 绑定在 character_designer 模块上。
    # 用 sys.modules[__name__] 兼容两种场景，确保 patch 生效。
    _current_module = sys.modules[__name__]

    mock_characters = [
        {
            "name": "林夏",
            "role": "主角",
            "description": "25岁女性，黑色短发，穿黑色风衣，眼神坚定",
        },
        {
            "name": "陈默",
            "role": "配角",
            "description": "30岁男性，戴金丝眼镜，穿灰色西装",
        },
    ]

    mock_scenes = [
        {
            "scene_id": 1,
            "scene_name": "雨夜街头",
            "description": "霓虹灯闪烁的东京街头，湿漉漉的地面反射着灯光",
        },
        {
            "scene_id": 2,
            "scene_name": "深夜便利店",
            "description": "冷白色荧光灯照亮的小便利店，货架整齐",
        },
    ]

    # Mock LLM 返回的角色卡 JSON
    mock_character_cards_json = json.dumps([
        {
            "name": "林夏",
            "role": "主角",
            "card_prompt": "年轻女性角色设定卡，三视图并排：正面+侧面+背面全身立绘，"
                           "25岁，黑色齐肩短发，穿长款黑色风衣，内搭白T恤和黑色紧身裤，"
                           "黑色短靴，眼神坚定，纯白色背景，无环境干扰，"
                           "电影感光影，35mm镜头质感",
            "portrait_prompt": "年轻女性单人正面半身像，25岁，黑色齐肩短发，"
                               "穿黑色风衣，眼神坚定，纯白色背景，自然光，"
                               "电影感光影，35mm镜头质感",
            "expression_prompts": {
                "正面平静": "年轻女性正面半身像，黑色短发，黑色风衣，平静表情，白背景",
                "微笑": "年轻女性正面半身像，黑色短发，黑色风衣，微笑，白背景",
                "严肃": "年轻女性正面半身像，黑色短发，黑色风衣，严肃表情，白背景",
                "惊讶": "年轻女性正面半身像，黑色短发，黑色风衣，惊讶表情，白背景",
            },
            "color_palette": "黑色+白色+暗银",
        },
        {
            "name": "陈默",
            "role": "配角",
            "card_prompt": "男性角色设定卡，三视图：正面+侧面+背面，"
                           "30岁，戴金丝眼镜，灰色西装三件套，白衬衫，"
                           "黑色皮鞋，纯白色背景，电影感光影，35mm镜头质感",
            "portrait_prompt": "男性单人正面半身像，30岁，金丝眼镜，灰色西装，白背景",
            "expression_prompts": {
                "正面平静": "男性正面半身像，金丝眼镜，灰色西装，平静，白背景",
                "微笑": "男性正面半身像，金丝眼镜，灰色西装，微笑，白背景",
                "严肃": "男性正面半身像，金丝眼镜，灰色西装，严肃，白背景",
                "惊讶": "男性正面半身像，金丝眼镜，灰色西装，惊讶，白背景",
            },
            "color_palette": "灰色+白色+金色",
        },
    ], ensure_ascii=False)

    mock_scene_cards_json = json.dumps([
        {
            "scene_id": 1,
            "scene_name": "雨夜街头",
            "grid_prompt": "九宫格场景设定卡，3x3 网格 9 个视角的同一空间："
                           "正面/左45/右45/左90/右90/背面/俯视/仰视/特写，"
                           "霓虹灯闪烁的东京街头，湿漉漉的地面反射灯光，"
                           "统一冷蓝调光影，无人物，电影感光影，35mm镜头质感",
            "wide_shot_prompt": "广角全景，东京街头夜景，霓虹灯，湿润地面反射，"
                                "无人物，冷蓝调，电影感光影，35mm镜头质感",
            "lighting_prompt": "冷蓝调主光+霓虹粉紫点缀，侧逆光，雨夜氛围",
            "color_palette": "深蓝+霓虹粉+紫红",
        },
        {
            "scene_id": 2,
            "scene_name": "深夜便利店",
            "grid_prompt": "九宫格场景设定卡，3x3 网格 9 个视角："
                           "正面/左45/右45/左90/右90/背面/俯视/仰视/特写，"
                           "冷白色荧光灯照亮的小便利店，货架整齐，"
                           "统一冷白光，无人物，电影感光影，35mm镜头质感",
            "wide_shot_prompt": "广角全景，深夜便利店内部，冷白荧光灯，整齐货架，"
                                "无人物，电影感光影，35mm镜头质感",
            "lighting_prompt": "冷白色荧光灯顶光，货架阴影，冰冷氛围",
            "color_palette": "冷白+浅蓝+灰色",
        },
    ], ensure_ascii=False)

    print("=" * 60)
    print("测试 1：generate_character_cards（mock LLM）")
    print("=" * 60)
    with patch.object(_current_module, "_chat", return_value=mock_character_cards_json):
        cards = generate_character_cards(mock_characters, art_style="cinematic")
    print(f"生成 {len(cards)} 张角色卡")
    print(f"第一张：{cards[0]['name']} - card_prompt 长度 {len(cards[0]['card_prompt'])}")
    assert cards[0]["color_palette"] == "黑色+白色+暗银", "mock 未生效，color_palette 不匹配"

    print()
    print("=" * 60)
    print("测试 2：generate_scene_cards（mock LLM）")
    print("=" * 60)
    with patch.object(_current_module, "_chat", return_value=mock_scene_cards_json):
        sc_cards = generate_scene_cards(mock_scenes, art_style="cinematic")
    print(f"生成 {len(sc_cards)} 张场景卡")
    print(f"第一张：场景{sc_cards[0]['scene_id']} {sc_cards[0]['scene_name']}")
    assert sc_cards[0]["color_palette"] == "深蓝+霓虹粉+紫红", "mock 未生效，color_palette 不匹配"

    print()
    print("=" * 60)
    print("测试 3：cards_to_markdown")
    print("=" * 60)
    md = cards_to_markdown(cards, sc_cards)
    print(md[:500] + "..." if len(md) > 500 else md)
    print(f"\n（Markdown 总长度：{len(md)} 字符）")

    print()
    print("=" * 60)
    print("测试 4：cards_to_docx")
    print("=" * 60)
    try:
        buf = cards_to_docx(cards, sc_cards, script_title="测试短片")
        data = buf.read()
        print(f"docx 生成成功，大小 {len(data)} 字节")
        # 写到临时文件验证
        out_path = "/data/user/work/test_character_cards.docx"
        with open(out_path, "wb") as f:
            f.write(data)
        print(f"已写入 {out_path}")
    except RuntimeError as e:
        print(f"跳过 docx 测试：{e}")

    print()
    print("所有测试通过 ✅")
